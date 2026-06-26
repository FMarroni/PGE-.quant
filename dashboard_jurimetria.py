"""
dashboard_jurimetria.py  –  v7
PGE .quant — Análise Quantitativa do Contencioso
Identidade Visual Institucional SP · Multi-Núcleo · Upload Integrado
Dois arquivos de entrada: Relatório de Processos + Relatório de Demandas
"""

import io
import json
import unicodedata
import urllib.request
from datetime import date, datetime
from pathlib import Path

import pandas as pd
import plotly.express as px
import streamlit as st

import db as _db

# =============================================================================
# CONFIGURAÇÃO DA PÁGINA  (obrigatoriamente a primeira chamada Streamlit)
# =============================================================================
st.set_page_config(
    layout="wide",
    page_title="PGE .quant",
    page_icon="⚖️",
    initial_sidebar_state="expanded",
)

# =============================================================================
# CSS — IDENTIDADE VISUAL SP  (Branco · Cinza Bandeirante · Vermelho)
# =============================================================================
_CSS = """
<style>
/* ── Reset e fundo branco ────────────────────────────────────────────── */
[data-testid="stAppViewContainer"] > .main  { background:#FFFFFF; }
[data-testid="stHeader"],
[data-testid="stDecoration"],
[data-testid="stToolbar"]                   { display:none !important; }
#MainMenu, footer                           { visibility:hidden; }
[data-testid="block-container"]             { padding-top: 1rem; }

/* ── Sidebar — Cinza Bandeirante ─────────────────────────────────────── */
[data-testid="stSidebar"] {
    background-color: #3D3D3D;
    border-right: 3px solid #CC0000;
}
[data-testid="stSidebar"] > div:first-child {
    background-color: #3D3D3D;
}

/* ── Todos os textos na sidebar: brancos ─────────────────────────────── */
[data-testid="stSidebar"] p,
[data-testid="stSidebar"] label,
[data-testid="stSidebar"] li,
[data-testid="stSidebar"] small,
[data-testid="stSidebar"] .stMarkdown,
[data-testid="stSidebar"] .stMarkdown *         { color:#E0E0E0 !important; }

[data-testid="stSidebar"] [data-testid="stWidgetLabel"],
[data-testid="stSidebar"] [data-testid="stWidgetLabel"] p,
[data-testid="stSidebar"] [data-testid="stWidgetLabel"] span { color:#E0E0E0 !important; }

[data-testid="stSidebar"] [data-baseweb="input"] input {
    background-color: #505050 !important;
    color: #FFFFFF !important;
    border-color: #666666 !important;
}
[data-testid="stSidebar"] [data-baseweb="input"] input:focus {
    border-color:#CC0000 !important;
    box-shadow: 0 0 0 2px rgba(204,0,0,.35) !important;
}

[data-testid="stSidebar"] [data-baseweb="select"] > div {
    background-color: #505050 !important;
    border-color: #666666 !important;
    color: #FFFFFF !important;
}
[data-testid="stSidebar"] [data-baseweb="select"] span,
[data-testid="stSidebar"] [data-baseweb="select"] div  { color:#FFFFFF !important; }
[data-testid="stSidebar"] [data-baseweb="select"] svg  { fill:#E0E0E0 !important; }

[data-testid="stSidebar"] [data-baseweb="tag"] {
    background-color:#CC0000 !important;
    color:#FFFFFF !important;
}
[data-testid="stSidebar"] [data-baseweb="tag"] span,
[data-testid="stSidebar"] [data-baseweb="tag"] * {
    color:#FFFFFF !important;
}

[data-testid="stSidebar"] .stButton button {
    background:#CC0000 !important;
    color:#FFFFFF !important;
    border:none !important;
    border-radius:5px !important;
    font-weight:600 !important;
    font-size:.8rem !important;
    transition: background .2s;
}
[data-testid="stSidebar"] .stButton button:hover {
    background:#A50000 !important;
}

[data-testid="stSidebar"] hr {
    border-color: rgba(255,255,255,.15) !important;
}

/* ── Cards de KPI ────────────────────────────────────────────────────── */
[data-testid="stMetric"] {
    background:    #FFFFFF;
    border:        1px solid #DEE2E6;
    border-top:    3px solid #CC0000;
    border-radius: 7px;
    padding:       0.85rem 1rem 0.7rem;
    box-shadow:    0 1px 4px rgba(0,0,0,0.05);
}
[data-testid="stMetricLabel"] > div {
    font-size:      .67rem !important;
    font-weight:    700 !important;
    text-transform: uppercase;
    letter-spacing: 1px;
    color:          #6C757D !important;
}
[data-testid="stMetricValue"] > div {
    font-size:   1.38rem !important;
    font-weight: 700 !important;
    color:       #212529 !important;
}
[data-testid="stMetricDelta"] svg { display:none; }
[data-testid="stMetricDelta"] > div {
    font-size:.71rem !important;
    color:#6C757D !important;
}

/* ── Abas ────────────────────────────────────────────────────────────── */
[data-testid="stTabs"] {
    border-bottom:2px solid #DEE2E6;
    margin-bottom:.4rem;
}
button[data-baseweb="tab"] {
    font-weight:600;
    font-size:.88rem;
    color:#6C757D;
    padding-bottom:.6rem;
}
button[aria-selected="true"] {
    color:#CC0000 !important;
    border-bottom:2px solid #CC0000 !important;
}

/* ── Rótulos de seção ────────────────────────────────────────────────── */
.sec-title {
    font-size:.62rem; font-weight:800; text-transform:uppercase;
    letter-spacing:2px; color:#CC0000; margin-bottom:.5rem;
}
.sec-sub {
    font-size:.76rem; color:#6C757D; margin-top:0; margin-bottom:1rem;
}

/* ── Barra vermelha superior ─────────────────────────────────────────── */
.top-bar {
    background:linear-gradient(100deg,#8B0000 0%,#CC0000 55%,#E63300 100%);
    padding:.9rem 1.5rem; border-radius:8px; margin-bottom:1rem;
    display:flex; align-items:center; justify-content:space-between; gap:1rem;
}
.top-bar h1  { color:#fff; font-size:1.35rem; font-weight:700; margin:0; letter-spacing:.4px; }
.top-bar p   { color:rgba(255,255,255,.78); font-size:.75rem; margin:2px 0 0; }
.top-badge {
    background:rgba(255,255,255,.14); border:1px solid rgba(255,255,255,.28);
    color:#fff; font-size:.7rem; padding:.28rem .85rem;
    border-radius:20px; white-space:nowrap; flex-shrink:0;
}

/* ── Divisor ─────────────────────────────────────────────────────────── */
.divider {
    height:2px;
    background:linear-gradient(90deg,#CC0000 0%,rgba(204,0,0,.08) 100%);
    border:none; margin:.85rem 0;
}

/* ── Área de upload ──────────────────────────────────────────────────── */
.upload-card {
    background:#F8F9FA; border:1px solid #DEE2E6; border-radius:10px;
    padding:1.5rem; margin-bottom:1rem;
}
</style>
"""
st.markdown(_CSS, unsafe_allow_html=True)

# =============================================================================
# CONSTANTES
# =============================================================================
DB_PATH     = _db.sqlite_path()   # referência ao caminho SQLite (display/backup)
GEOJSON_URL = (
    "https://raw.githubusercontent.com/tbrugz/geodata-br"
    "/master/geojson/geojs-35-mun.json"
)

_STATUS_CORES = {
    "Vitória":      "#27AE60",
    "Perda":        "#CC0000",
    "Em Andamento": "#ADB5BD",
}
_ESCALA_DIVERGENTE = [
    [0.00, "#CC0000"],
    [0.25, "#E07B00"],
    [0.50, "#F5C518"],
    [0.75, "#7DC95E"],
    [1.00, "#1A8A2E"],
]
_PREFIXO_ASSUNTO = "Sistema Remuneratório: Verbas - "

_COLUNAS_TABELA = {
    "processo":             "Nº do Processo",
    "valor":                "Valor (R$)",
    "assunto_label":        "Assunto",
    "procurador":           "Procurador",
    "vara":                 "Vara",
    "ult_demanda":          "Última Demanda",
    "data_ultima_demanda":  "Data",
    "status_exito":         "Status",
}

# ── Núcleos da Coordenadoria do Contencioso Geral ─────────────────────────
_NUCLEOS = [
    "Núcleo Imobiliário",
    "Núcleo Ambiental",
    "Núcleo de Pessoal Residual",
    "Núcleo de Pessoal Educação",
    "Núcleo de Pessoal das Carreiras de Estado",
    "Núcleo de Pessoal Militar",
    "Núcleo Previdenciário",
    "Núcleo de Gestão e Prevenção de Demandas Repetitivas (GPDR)",
    "Núcleo Trabalhista",
    "Núcleo Estratégico de Pessoal e Previdenciário (NEPP)",
    "Núcleo de Responsabilidade Subsidiária Trabalhista (NRST)",
    "Núcleo de Responsabilidade Civil",
    "Núcleo de Contencioso PGE/DETRAN-SP",
    "Núcleo de Saúde Pública",
    "Núcleo de Regulação e Contratações Públicas",
    "Núcleo de Poder de Polícia",
    "Núcleo de Políticas Públicas",
    "Núcleo de Propositura e Acompanhamento (NPA)",
    "Coordenadoria de Execuções contra a Fazenda Pública (CEFAP)",
]

# ── Mapeamento de colunas — Arquivo A (Processos) ────────────────────────
# Colunas Últ. Andamento Judicial e Data do Andamento removidas da fonte;
# mantidas no schema do banco para preservar dados históricos.
_MAPA_COLUNAS: dict[str, str] = {
    "Pasta":                              "pasta",
    "Processo":                           "processo",
    "Valor":                              "valor",
    "Ajuizamento":                        "ajuizamento",
    "Cadastro":                           "cadastro",
    "Classe":                             "classe",
    "Matéria":                            "materia",
    "Assuntos":                           "assuntos",
    "Tribunal":                           "tribunal",
    "Unidade Judicial / Instituição":     "unidade_judicial",
    "Vara":                               "vara",
    "Polo PGE":                           "polo_pge",
    "Qualificação":                       "qualificacao",
    "Parte Representada":                 "parte_representada",
    "Documento Parte Rep.":               "documento_parte_rep",
    "Parte Contrária":                    "parte_contraria",
    "Documento":                          "documento",
    "Outras Partes Ativas":               "outras_partes_ativas",
    "Outras Partes Passivas":             "outras_partes_passivas",
    "Advogados":                          "advogados",
    "Tramitação":                         "tramitacao",
    "Situação":                           "situacao",
    "Unidade":                            "unidade",
    "Mesa":                               "mesa",
    "Tipo Distribuição":                  "tipo_distribuicao",
    "Nº de Dívidas":                      "num_dividas",
    "Soma Val. Atualizados":              "soma_val_atualizados",
}

# ── Mapeamento de colunas — Arquivo B (Demandas) ─────────────────────────
_MAPA_COLUNAS_DEMANDAS: dict[str, str] = {
    "Unidade":      "unidade",
    "Procurador":   "procurador",
    "Demanda":      "demanda",
    "Processo":     "processo_orig",
    "Pasta":        "pasta",
    "Qualificação": "qualificacao",
    "Matéria":      "materia",
    "Assuntos":     "assuntos",
    "Tribunal":     "tribunal",
    "Origem":       "origem",
    "Status":       "status_demanda",
    "Entrada":      "entrada",
    "Conclusão":    "conclusao",
    "Horas":        "horas",
    "Publicou":     "publicou",
}

# ── Schema DDL ────────────────────────────────────────────────────────────
_DDL_PROCESSOS = """
CREATE TABLE IF NOT EXISTS processos_consolidados (
    processo                    TEXT PRIMARY KEY,
    pasta                       TEXT,
    valor                       REAL,
    ajuizamento                 TEXT,
    cadastro                    TEXT,
    classe                      TEXT,
    materia                     TEXT,
    assuntos                    TEXT,
    tribunal                    TEXT,
    unidade_judicial             TEXT,
    vara                        TEXT,
    polo_pge                    TEXT,
    qualificacao                TEXT,
    parte_representada          TEXT,
    documento_parte_rep         TEXT,
    parte_contraria             TEXT,
    documento                   TEXT,
    outras_partes_ativas        TEXT,
    outras_partes_passivas      TEXT,
    advogados                   TEXT,
    ult_andamento_judicial       TEXT,
    data_do_andamento           TEXT,
    tramitacao                  TEXT,
    situacao                    TEXT,
    unidade                     TEXT,
    mesa                        TEXT,
    tipo_distribuicao           TEXT,
    num_dividas                 TEXT,
    soma_val_atualizados        TEXT,
    status_exito                TEXT,
    nucleo                      TEXT,
    data_ultima_atualizacao     TEXT,
    procurador                  TEXT,
    ult_demanda                 TEXT,
    data_ultima_demanda         TEXT,
    total_horas                 REAL
)
"""

_DDL_DEMANDAS = """
CREATE TABLE IF NOT EXISTS demandas (
    id              INTEGER PRIMARY KEY AUTOINCREMENT,
    processo_base   TEXT NOT NULL,
    processo_orig   TEXT,
    pasta           TEXT,
    unidade         TEXT,
    procurador      TEXT,
    demanda         TEXT,
    qualificacao    TEXT,
    materia         TEXT,
    assuntos        TEXT,
    tribunal        TEXT,
    origem          TEXT,
    status_demanda  TEXT,
    entrada         TEXT,
    conclusao       TEXT,
    horas           REAL,
    publicou        TEXT,
    nucleo          TEXT,
    data_upload     TEXT
)
"""

_DDL_CONTROLE = """
CREATE TABLE IF NOT EXISTS controle_uploads (
    id                              INTEGER PRIMARY KEY AUTOINCREMENT,
    data_upload                     TEXT    NOT NULL,
    nome_arquivo                    TEXT    NOT NULL,
    quantidade_registros_processados INTEGER NOT NULL,
    nucleo                          TEXT
)
"""

_COLUNAS_SCHEMA_DB = (
    "processo", "pasta", "valor", "ajuizamento", "cadastro", "classe",
    "materia", "assuntos", "tribunal", "unidade_judicial", "vara",
    "polo_pge", "qualificacao", "parte_representada", "documento_parte_rep",
    "parte_contraria", "documento", "outras_partes_ativas",
    "outras_partes_passivas", "advogados", "ult_andamento_judicial",
    "data_do_andamento", "tramitacao", "situacao", "unidade", "mesa",
    "tipo_distribuicao", "num_dividas", "soma_val_atualizados",
    "status_exito", "nucleo", "data_ultima_atualizacao",
    "procurador", "ult_demanda", "data_ultima_demanda", "total_horas",
)

# status_exito não atualizado pelo UPSERT do Arquivo A; é derivado do Arquivo B
_COLUNAS_ATUALIZAVEIS = (
    "valor",
    "tramitacao",
    "situacao",
    "nucleo",
    "data_ultima_atualizacao",
)

_COLUNAS_DEMANDAS_DB = (
    "processo_base", "processo_orig", "pasta", "unidade", "procurador",
    "demanda", "qualificacao", "materia", "assuntos", "tribunal",
    "origem", "status_demanda", "entrada", "conclusao", "horas",
    "publicou", "nucleo", "data_upload",
)

_REGRAS_EXITO: list[tuple[str, str]] = [
    ("Favorável ao Estado",               "Vitória"),
    ("Extinção sem julgamento do mérito", "Vitória"),
    ("Desfavorável ao Estado",            "Perda"),
]

# =============================================================================
# MIGRAÇÃO DE BANCO
# =============================================================================

def _migrar_banco() -> None:
    """Migração não-destrutiva: adiciona colunas novas e cria tabela demandas."""
    if not _db.db_exists():
        return
    con = _db.connect()
    try:
        cur = con.cursor()
        for col, tipo in [
            ("nucleo",               "TEXT"),
            ("procurador",           "TEXT"),
            ("ult_demanda",          "TEXT"),
            ("data_ultima_demanda",  "TEXT"),
            ("total_horas",          "REAL"),
        ]:
            _db.add_column_if_not_exists(cur, "processos_consolidados", col, tipo)
        cur.execute(_db.adapt_sql(_DDL_DEMANDAS))
        cur.execute(_db.adapt_sql(_DDL_CONTROLE))
        _db.add_column_if_not_exists(cur, "controle_uploads", "nucleo", "TEXT")
        con.commit()
    finally:
        con.close()

# =============================================================================
# UTILITÁRIOS
# =============================================================================

def _sem_acento(texto: str) -> str:
    nfd = unicodedata.normalize("NFD", str(texto))
    return "".join(c for c in nfd if not unicodedata.combining(c))


# ── Normalização para detecção automática de núcleo ───────────────────────
_NUCLEOS_NORM: dict[str, str] = {
    _sem_acento(n).upper().strip(): n for n in _NUCLEOS
}

# Preposições/artigos que podem variar entre o nome oficial cadastrado e o
# texto exportado pelo sistema de origem (ex.: "Pessoal Carreiras" vs.
# "Pessoal das Carreiras") sem mudar o núcleo referido — ignorados na
# comparação de fallback para tolerar essas pequenas divergências de texto.
_STOPWORDS_NUCLEO = {"DE", "DA", "DO", "DAS", "DOS", "E"}


def _normaliza_nucleo_sem_stopwords(s: str) -> str:
    palavras = _sem_acento(s).upper().split()
    return " ".join(p for p in palavras if p not in _STOPWORDS_NUCLEO)


_NUCLEOS_NORM_SEM_STOPWORDS: dict[str, str] = {
    _normaliza_nucleo_sem_stopwords(n): n for n in _NUCLEOS
}


def _detectar_nucleo_label(label: str) -> str | None:
    """Mapeia um valor bruto da coluna 'Unidade' (ex.: 'Núcleo de Pessoal Militar - NPM')
    para o nome canônico em _NUCLEOS, tolerando acentos, caixa, sufixo de sigla e
    pequenas divergências de artigos/preposições (ex.: 'Pessoal Carreiras' vs.
    'Pessoal das Carreiras')."""
    if not label or pd.isna(label):
        return None
    bruto = str(label).strip()

    # tenta o valor completo primeiro (núcleos cuja sigla já é parte do nome oficial)
    direto = _NUCLEOS_NORM.get(_sem_acento(bruto).upper().strip())
    if direto:
        return direto

    # remove sufixo " - SIGLA" e tenta novamente
    sem_sigla = bruto.split(" - ")[0].strip()
    direto = _NUCLEOS_NORM.get(_sem_acento(sem_sigla).upper().strip())
    if direto:
        return direto

    # fallback tolerante a artigos/preposições divergentes
    return _NUCLEOS_NORM_SEM_STOPWORDS.get(_normaliza_nucleo_sem_stopwords(sem_sigla))


def _ler_csv_upload(conteudo: bytes) -> pd.DataFrame:
    """Lê CSV de upload com fallback de encoding. Usado tanto na detecção
    automática de núcleo quanto no processamento completo do arquivo."""
    for enc in ("utf-8-sig", "utf-8", "latin-1", "cp1252"):
        try:
            return pd.read_csv(
                io.BytesIO(conteudo),
                sep=",",
                quotechar='"',
                engine="python",
                encoding=enc,
                dtype=str,
                keep_default_na=False,
                na_values=[],
            )
        except (UnicodeDecodeError, Exception):
            continue
    raise ValueError(
        "Não foi possível decodificar o arquivo. "
        "Certifique-se de exportar como UTF-8 ou Latin-1."
    )


def _detectar_nucleo_arquivo(arquivo) -> str | None:
    """Lê a coluna 'Unidade' de um arquivo de upload e identifica o núcleo
    mais frequente, sem consumir o ponteiro do arquivo (preserva para reuso)."""
    if arquivo is None:
        return None
    try:
        arquivo.seek(0)
        conteudo = arquivo.read()
        arquivo.seek(0)
    except Exception:
        return None

    try:
        df_amostra = _ler_csv_upload(conteudo)
    except Exception:
        return None

    col_unidade = next(
        (c for c in df_amostra.columns if str(c).strip() == "Unidade"), None
    )
    if col_unidade is None:
        return None

    valores = df_amostra[col_unidade].astype(str).str.strip()
    valores = valores[valores != ""]
    if valores.empty:
        return None

    mais_frequente = valores.mode().iloc[0]
    return _detectar_nucleo_label(mais_frequente)


def _grupo_do_nucleo(nucleo: str, grupo_opts: dict[str, list[str]]) -> str | None:
    for grupo, lista in grupo_opts.items():
        if nucleo in lista:
            return grupo
    return None


def _brl(valor: float) -> str:
    return "R$ " + f"{valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def _kpis(df: pd.DataFrame) -> dict:
    v   = (df["status_exito"] == "Vitória").sum()
    p   = (df["status_exito"] == "Perda").sum()
    dec = v + p
    return {
        "total":     len(df),
        "em_risco":  df.loc[df["status_exito"] == "Em Andamento", "valor"].sum(),
        "economia":  df.loc[df["status_exito"] == "Vitória",      "valor"].sum(),
        "taxa":      (v / dec * 100) if dec else 0.0,
        "vitorias":  int(v),
        "perdas":    int(p),
        "decididos": int(dec),
    }

# =============================================================================
# MOTOR DE INGESTÃO  (dois arquivos)
# =============================================================================

def processar_upload(
    arquivo_processos,
    arquivo_demandas,
    nome_nucleo: str,
) -> tuple[int, int]:
    """
    Processa Arquivo A (processos) e Arquivo B (demandas), persiste no SQLite
    e deriva campos cruzados. Retorna (qtd_processos, qtd_demandas).
    """

    # ── 1. Arquivo A — Processos ───────────────────────────────────────────
    df = _ler_csv_upload(arquivo_processos.read())

    colunas_validas = [
        c for c in df.columns
        if str(c).strip() not in ("", "R$") and not str(c).startswith("Unnamed:")
    ]
    df = df[colunas_validas].copy()
    df.columns = [c.strip() for c in df.columns]
    df = df.rename(columns=_MAPA_COLUNAS)

    for col in ("ajuizamento", "cadastro"):
        if col in df.columns:
            df[col] = pd.to_datetime(df[col], format="%d/%m/%Y", errors="coerce")

    if "valor" not in df.columns:
        raise KeyError(
            "Coluna 'Valor' não encontrada no arquivo de processos. "
            "Verifique o formato do relatório."
        )
    df["valor"] = (
        df["valor"].str.strip()
        .str.replace(".", "", regex=False)
        .str.replace(",", ".", regex=False)
        .pipe(pd.to_numeric, errors="coerce")
    )

    if "processo" not in df.columns:
        raise KeyError("Coluna 'Processo' não encontrada no arquivo de processos.")
    df["processo"] = df["processo"].str.strip()
    colunas_cat = [c for c in df.columns if c not in ("processo", "valor")]
    agg: dict = {"valor": "sum"}
    agg.update({c: "first" for c in colunas_cat})
    df = df.groupby("processo", as_index=False, sort=False).agg(agg)

    df["nucleo"] = nome_nucleo
    df["status_exito"] = "Em Andamento"
    df["data_ultima_atualizacao"] = datetime.now().isoformat(timespec="seconds")

    for col in ("ajuizamento", "cadastro"):
        if col in df.columns and pd.api.types.is_datetime64_any_dtype(df[col]):
            df[col] = df[col].dt.strftime("%Y-%m-%d")

    df = df.where(pd.notnull(df), other=None)

    colunas_banco = [c for c in _COLUNAS_SCHEMA_DB if c in df.columns]
    df_banco = df[colunas_banco]
    cols_str   = ", ".join(colunas_banco)
    placeh_str = ", ".join([_db.ph()] * len(colunas_banco))
    atualizacoes = ", ".join(
        f"{c} = excluded.{c}"
        for c in _COLUNAS_ATUALIZAVEIS
        if c in colunas_banco
    )
    sql_upsert = f"""
        INSERT INTO processos_consolidados ({cols_str})
        VALUES ({placeh_str})
        ON CONFLICT(processo) DO UPDATE SET
            {atualizacoes}
    """
    registros_proc = [tuple(row) for row in df_banco.itertuples(index=False, name=None)]

    # ── 2. Arquivo B — Demandas ────────────────────────────────────────────
    df_dem = _ler_csv_upload(arquivo_demandas.read())

    colunas_validas_d = [
        c for c in df_dem.columns
        if str(c).strip() not in ("", "R$") and not str(c).startswith("Unnamed:")
    ]
    df_dem = df_dem[colunas_validas_d].copy()
    df_dem.columns = [c.strip() for c in df_dem.columns]
    df_dem = df_dem.rename(columns=_MAPA_COLUNAS_DEMANDAS)

    if "processo_orig" not in df_dem.columns:
        raise KeyError("Coluna 'Processo' não encontrada no arquivo de demandas.")

    df_dem["processo_orig"] = df_dem["processo_orig"].str.strip()
    df_dem["processo_base"] = df_dem["processo_orig"].str.split("/").str[0]

    for col in ("entrada", "conclusao"):
        if col in df_dem.columns:
            df_dem[col] = pd.to_datetime(df_dem[col], format="%d/%m/%Y", errors="coerce")

    if "horas" in df_dem.columns:
        df_dem["horas"] = (
            df_dem["horas"].str.strip()
            .str.replace(",", ".", regex=False)
            .pipe(pd.to_numeric, errors="coerce")
            .fillna(0.0)
        )

    df_dem["nucleo"] = nome_nucleo
    ts = datetime.now().isoformat(timespec="seconds")
    df_dem["data_upload"] = ts

    # ── 3. Derivação cruzada ───────────────────────────────────────────────
    df_sorted = df_dem.sort_values("conclusao", ascending=False, na_position="last")

    mais_recente = df_sorted.groupby("processo_base", sort=False).first().reset_index()

    if "horas" in df_dem.columns:
        total_horas = (
            df_dem.groupby("processo_base")["horas"]
            .sum().reset_index(name="total_horas")
        )
    else:
        total_horas = pd.DataFrame(columns=["processo_base", "total_horas"])

    def _classif_exito_serie(series_list: list) -> str:
        for dem in series_list:
            if not dem or pd.isna(dem):
                continue
            for frag, status in _REGRAS_EXITO:
                if frag in str(dem):
                    return status
        return "Em Andamento"

    status_por_proc = (
        df_sorted.groupby("processo_base", sort=False)["demanda"]
        .apply(lambda s: _classif_exito_serie(s.tolist()))
        .reset_index(name="status_exito")
    ) if "demanda" in df_sorted.columns else pd.DataFrame(
        columns=["processo_base", "status_exito"]
    )

    df_deriv = mais_recente[["processo_base"]].copy()
    df_deriv["ult_demanda"] = (
        mais_recente["demanda"].values if "demanda" in mais_recente.columns else None
    )
    df_deriv["data_ultima_demanda"] = (
        mais_recente["conclusao"].dt.strftime("%Y-%m-%d").values
        if "conclusao" in mais_recente.columns
        else None
    )
    df_deriv["procurador"] = (
        mais_recente["procurador"].values if "procurador" in mais_recente.columns else None
    )
    df_deriv = df_deriv.merge(total_horas, on="processo_base", how="left")
    df_deriv = df_deriv.merge(status_por_proc, on="processo_base", how="left")
    df_deriv["total_horas"]  = df_deriv.get("total_horas",  pd.Series(dtype=float)).fillna(0.0)
    df_deriv["status_exito"] = df_deriv.get("status_exito", pd.Series(dtype=str)).fillna("Em Andamento")
    df_deriv = df_deriv.where(pd.notnull(df_deriv), other=None)

    update_records = [
        (
            row["procurador"],
            row["ult_demanda"],
            row["data_ultima_demanda"],
            row.get("total_horas"),
            row["status_exito"],
            ts,
            row["processo_base"],
        )
        for _, row in df_deriv.iterrows()
    ]

    # ── 4. Preparar demandas para storage (datas → string) ─────────────────
    for col in ("entrada", "conclusao"):
        if col in df_dem.columns and pd.api.types.is_datetime64_any_dtype(df_dem[col]):
            df_dem[col] = df_dem[col].dt.strftime("%Y-%m-%d")

    df_dem = df_dem.where(pd.notnull(df_dem), other=None)
    cols_dem   = [c for c in _COLUNAS_DEMANDAS_DB if c in df_dem.columns]
    df_dem_db  = df_dem[cols_dem]
    placeh_dem = ", ".join([_db.ph()] * len(cols_dem))
    cols_dem_str = ", ".join(cols_dem)
    sql_insert_dem = f"INSERT INTO demandas ({cols_dem_str}) VALUES ({placeh_dem})"
    registros_dem = [tuple(row) for row in df_dem_db.itertuples(index=False, name=None)]

    _p = _db.ph()
    sql_update_deriv = (
        f"UPDATE processos_consolidados "
        f"SET procurador = {_p}, ult_demanda = {_p}, data_ultima_demanda = {_p}, "
        f"    total_horas = {_p}, status_exito = {_p}, data_ultima_atualizacao = {_p} "
        f"WHERE processo = {_p}"
    )
    sql_del_dem      = _db.adapt_sql("DELETE FROM demandas WHERE nucleo = ?")
    sql_ins_ctrl     = _db.adapt_sql(
        "INSERT INTO controle_uploads "
        "(data_upload, nome_arquivo, quantidade_registros_processados, nucleo) VALUES (?, ?, ?, ?)"
    )

    # ── 5. Persistir no banco (transação única) ────────────────────────────
    con = _db.connect()
    try:
        cur = con.cursor()
        cur.execute(_db.adapt_sql(_DDL_PROCESSOS))
        cur.execute(_db.adapt_sql(_DDL_DEMANDAS))
        cur.execute(_db.adapt_sql(_DDL_CONTROLE))
        for col, tipo in [
            ("nucleo",              "TEXT"),
            ("procurador",          "TEXT"),
            ("ult_demanda",         "TEXT"),
            ("data_ultima_demanda", "TEXT"),
            ("total_horas",         "REAL"),
        ]:
            _db.add_column_if_not_exists(cur, "processos_consolidados", col, tipo)
        _db.add_column_if_not_exists(cur, "controle_uploads", "nucleo", "TEXT")

        _db.executemany(cur, sql_upsert, registros_proc)
        cur.execute(sql_del_dem, (nome_nucleo,))
        _db.executemany(cur, sql_insert_dem, registros_dem)
        _db.executemany(cur, sql_update_deriv, update_records)
        cur.execute(sql_ins_ctrl, (ts, arquivo_processos.name, len(registros_proc), nome_nucleo))
        cur.execute(sql_ins_ctrl, (ts, arquivo_demandas.name, len(registros_dem), nome_nucleo))
        con.commit()
    except Exception:
        con.rollback()
        raise
    finally:
        con.close()

    return len(registros_proc), len(registros_dem)

# =============================================================================
# EXCLUSÃO DE DADOS
# =============================================================================

def _nucleos_carregados() -> list[str]:
    """Lista núcleos com dados atualmente persistidos no banco."""
    if not _db.db_exists():
        return []
    try:
        con = _db.connect()
        df = _db.read_sql(
            "SELECT DISTINCT nucleo FROM processos_consolidados "
            "WHERE nucleo IS NOT NULL AND nucleo <> ''",
            con,
        )
        con.close()
    except Exception:
        return []
    return sorted(df["nucleo"].dropna().tolist())


def _resumo_geral() -> dict:
    """Conta processos, demandas e uploads no banco, independente de núcleo estar marcado."""
    resumo = {"processos": 0, "demandas": 0, "uploads": 0}
    if not _db.db_exists():
        return resumo
    try:
        con = _db.connect()
        cur = con.cursor()
        cur.execute("SELECT COUNT(*) FROM processos_consolidados")
        resumo["processos"] = cur.fetchone()[0]
        cur.execute("SELECT COUNT(*) FROM demandas")
        resumo["demandas"] = cur.fetchone()[0]
        cur.execute("SELECT COUNT(*) FROM controle_uploads")
        resumo["uploads"] = cur.fetchone()[0]
        con.close()
    except Exception:
        pass
    return resumo


def _resumo_nucleo(nucleo: str) -> dict:
    """Conta processos, demandas e uploads registrados para um núcleo."""
    resumo = {"processos": 0, "demandas": 0, "uploads": 0}
    if not _db.db_exists():
        return resumo
    try:
        con = _db.connect()
        cur = con.cursor()
        ph = _db.ph()
        cur.execute(f"SELECT COUNT(*) FROM processos_consolidados WHERE nucleo = {ph}", (nucleo,))
        resumo["processos"] = cur.fetchone()[0]
        cur.execute(f"SELECT COUNT(*) FROM demandas WHERE nucleo = {ph}", (nucleo,))
        resumo["demandas"] = cur.fetchone()[0]
        cur.execute(f"SELECT COUNT(*) FROM controle_uploads WHERE nucleo = {ph}", (nucleo,))
        resumo["uploads"] = cur.fetchone()[0]
        con.close()
    except Exception:
        pass
    return resumo


def excluir_dados_nucleo(nucleo: str) -> dict:
    """
    Exclui permanentemente todos os processos, demandas e registros de upload
    associados a um núcleo. Operação destrutiva e irreversível.
    Retorna dict com a contagem de linhas excluídas por tabela.
    """
    ph = _db.ph()
    con = _db.connect()
    try:
        cur = con.cursor()
        cur.execute(f"SELECT COUNT(*) FROM processos_consolidados WHERE nucleo = {ph}", (nucleo,))
        n_proc = cur.fetchone()[0]
        cur.execute(f"SELECT COUNT(*) FROM demandas WHERE nucleo = {ph}", (nucleo,))
        n_dem = cur.fetchone()[0]
        cur.execute(f"SELECT COUNT(*) FROM controle_uploads WHERE nucleo = {ph}", (nucleo,))
        n_ctrl = cur.fetchone()[0]

        cur.execute(f"DELETE FROM processos_consolidados WHERE nucleo = {ph}", (nucleo,))
        cur.execute(f"DELETE FROM demandas WHERE nucleo = {ph}", (nucleo,))
        cur.execute(f"DELETE FROM controle_uploads WHERE nucleo = {ph}", (nucleo,))
        con.commit()
    except Exception:
        con.rollback()
        raise
    finally:
        con.close()

    return {"processos": n_proc, "demandas": n_dem, "uploads": n_ctrl}


def excluir_todos_dados() -> dict:
    """
    Reset total: exclui permanentemente TODOS os processos, demandas e
    registros de upload de TODOS os núcleos. Operação destrutiva e irreversível.
    Retorna dict com a contagem de linhas excluídas por tabela.
    """
    con = _db.connect()
    try:
        cur = con.cursor()
        cur.execute("SELECT COUNT(*) FROM processos_consolidados")
        n_proc = cur.fetchone()[0]
        cur.execute("SELECT COUNT(*) FROM demandas")
        n_dem = cur.fetchone()[0]
        cur.execute("SELECT COUNT(*) FROM controle_uploads")
        n_ctrl = cur.fetchone()[0]

        cur.execute("DELETE FROM processos_consolidados")
        cur.execute("DELETE FROM demandas")
        cur.execute("DELETE FROM controle_uploads")
        con.commit()
    except Exception:
        con.rollback()
        raise
    finally:
        con.close()

    return {"processos": n_proc, "demandas": n_dem, "uploads": n_ctrl}

# =============================================================================
# CARREGAMENTO DE DADOS
# =============================================================================

@st.cache_data(ttl=300, show_spinner="⏳ Carregando base de processos...")
def _carregar_df() -> pd.DataFrame:
    con = _db.connect()
    df  = _db.read_sql(
        """
        SELECT processo, valor, ajuizamento, assuntos, tribunal,
               unidade_judicial, vara, situacao, unidade, mesa,
               status_exito, nucleo,
               procurador, ult_demanda, data_ultima_demanda, total_horas
        FROM   processos_consolidados
        """,
        con,
    )
    con.close()

    df["valor"]       = pd.to_numeric(df["valor"], errors="coerce").fillna(0.0)
    df["ajuizamento"] = pd.to_datetime(df["ajuizamento"], errors="coerce")
    df["nucleo"]      = df["nucleo"].fillna("(sem núcleo)")
    df["procurador"]  = df["procurador"].fillna("(sem procurador)")
    df["total_horas"] = pd.to_numeric(df["total_horas"], errors="coerce").fillna(0.0)

    df["data_ultima_demanda"] = pd.to_datetime(df["data_ultima_demanda"], errors="coerce")

    df["assunto_label"] = (
        df["assuntos"]
        .str.replace(_PREFIXO_ASSUNTO, "", regex=False)
        .str.split(",").str[0]
        .str.strip()
        .str[:60]
    )

    def _comarca(s: str) -> str:
        if not s or pd.isna(s):
            return ""
        s = str(s).strip()
        if s.startswith("Comarca de "):
            s = s[len("Comarca de "):]
        s = s.split(" - ")[0].strip()
        return _sem_acento(s).upper()

    df["comarca_limpa"] = df["unidade_judicial"].apply(_comarca)
    return df


@st.cache_data(ttl=300, show_spinner="⏳ Carregando demandas...")
def _carregar_demandas() -> pd.DataFrame:
    if not _db.db_exists():
        return pd.DataFrame()
    try:
        con = _db.connect()
        df = _db.read_sql(
            """SELECT processo_base, processo_orig, procurador, demanda,
                      unidade, origem, status_demanda,
                      entrada, conclusao, horas, publicou, nucleo
               FROM demandas""",
            con,
        )
        con.close()
    except Exception:
        return pd.DataFrame()

    df["conclusao"] = pd.to_datetime(df["conclusao"], errors="coerce")
    df["entrada"]   = pd.to_datetime(df["entrada"],   errors="coerce")
    df["horas"]     = pd.to_numeric(df["horas"], errors="coerce").fillna(0.0)
    return df


@st.cache_data(ttl=300)
def _ultimo_upload() -> str:
    if not _db.db_exists():
        return "sem registro"
    try:
        con = _db.connect()
        row = _db.read_sql(
            "SELECT data_upload FROM controle_uploads ORDER BY id DESC LIMIT 1", con
        )
        con.close()
    except Exception:
        return "sem registro"
    if row.empty:
        return "sem registro"
    return str(row["data_upload"].iloc[0])[:16].replace("T", " ")


@st.cache_data(ttl=86_400, show_spinner="🌍 Carregando mapa de SP...")
def _carregar_geojson() -> dict | None:
    try:
        req = urllib.request.Request(
            GEOJSON_URL,
            headers={"User-Agent": "jurimetria-pge/1.0"},
        )
        with urllib.request.urlopen(req, timeout=12) as resp:
            geo = json.loads(resp.read().decode("utf-8"))
        for feat in geo["features"]:
            nome = feat["properties"].get("name", "")
            feat["properties"]["nome_norm"] = _sem_acento(nome).upper()
        return geo
    except Exception:
        return None

# =============================================================================
# SIDEBAR  →  retorna DataFrame filtrado
# =============================================================================

def _sidebar(df_full: pd.DataFrame) -> pd.DataFrame:
    with st.sidebar:

        # ── Cabeçalho ─────────────────────────────────────────────────────
        _brasao = Path(__file__).parent / "__rectSitelogo__Brasão_do_estado_de_São_Paulo.svg-min.jpg"
        _, col_c, _ = st.columns([1, 2, 1])
        with col_c:
            if _brasao.exists():
                st.image(str(_brasao), width=120)
        st.markdown(
            "<div style='text-align:center;padding:.1rem 0 .9rem'>"
            "  <p style='color:#FFFFFF;font-size:1.05rem;font-weight:800;"
            "  letter-spacing:.5px;margin:0;line-height:1.2'>PGE <span style=\"color:#CC0000\">.quant</span></p>"
            "  <p style='color:#BDBDBD;font-size:.58rem;text-transform:uppercase;"
            "  letter-spacing:2px;margin:.3rem 0 0;font-weight:600'>Contencioso Geral · SP</p>"
            "</div>",
            unsafe_allow_html=True,
        )
        st.markdown(
            "<hr style='border:none;border-top:1px solid rgba(255,255,255,.15);margin:.2rem 0 .9rem'>",
            unsafe_allow_html=True,
        )
        st.markdown(
            "<p style='font-size:.62rem;text-transform:uppercase;letter-spacing:1.5px;"
            "color:#CC0000;margin-bottom:.5rem;font-weight:800'>🔍 Filtros</p>",
            unsafe_allow_html=True,
        )

        # ── Filtro 1 — Núcleos ─────────────────────────────────────────────
        nucleos_disponiveis = sorted(df_full["nucleo"].dropna().unique().tolist())
        nucleo_sel = st.multiselect(
            "Núcleos",
            options=nucleos_disponiveis,
            default=[],
            placeholder="Todos os núcleos",
            key="f_nucleo",
        )

        # ── Filtro 2 — Procurador ──────────────────────────────────────────
        proc_opts = sorted([
            p for p in df_full["procurador"].dropna().unique()
            if p != "(sem procurador)"
        ])
        procurador_sel = st.multiselect(
            "Procurador",
            options=proc_opts,
            default=[],
            placeholder="Todos os procuradores",
            key="f_procurador",
        )

        # ── Filtro 3 — Período (Última Demanda) ────────────────────────────
        datas = df_full["data_ultima_demanda"].dropna()
        d_min = datas.min().date() if not datas.empty else date(2024, 1, 1)
        d_max = datas.max().date() if not datas.empty else date.today()

        periodo = st.date_input(
            "Período (Última Demanda)",
            value=(d_min, d_max),
            min_value=d_min,
            max_value=d_max,
            format="DD/MM/YYYY",
            key="f_periodo",
        )

        # ── Filtro 4 — Status de Êxito ─────────────────────────────────────
        status_opts = sorted(df_full["status_exito"].dropna().unique().tolist())
        status_sel  = st.multiselect(
            "Status de Êxito",
            options=status_opts,
            default=status_opts,
            key="f_status",
        )

        # ── Filtro 5 — Assunto ─────────────────────────────────────────────
        labels_assunto = sorted(df_full["assunto_label"].dropna().unique().tolist())
        assunto_sel    = st.multiselect(
            "Assunto",
            options=labels_assunto,
            default=[],
            placeholder="Todos os assuntos",
            key="f_assunto",
        )

        # ── Limpar ─────────────────────────────────────────────────────────
        st.markdown("<div style='margin-top:.9rem'>", unsafe_allow_html=True)
        if st.button("✕  Limpar Filtros", use_container_width=True):
            for k in ("f_nucleo", "f_procurador", "f_periodo", "f_status", "f_assunto"):
                st.session_state.pop(k, None)
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

        # ── Rodapé ─────────────────────────────────────────────────────────
        st.markdown(
            f"<hr style='border:none;border-top:1px solid rgba(255,255,255,.15);margin:1rem 0 .5rem'>"
            f"<small style='color:#9E9E9E;font-size:.62rem'>"
            f"🕐 Última atualização<br>"
            f"<span style='color:#BDBDBD'>{_ultimo_upload()}</span>"
            f"</small>",
            unsafe_allow_html=True,
        )

    # ── Aplicar filtros ─────────────────────────────────────────────────────
    mask = pd.Series(True, index=df_full.index)

    if nucleo_sel:
        mask &= df_full["nucleo"].isin(nucleo_sel)

    if procurador_sel:
        mask &= df_full["procurador"].isin(procurador_sel)

    if isinstance(periodo, (list, tuple)) and len(periodo) == 2:
        d_ini, d_fim = periodo
        mask &= (
            df_full["data_ultima_demanda"].isna()
            | df_full["data_ultima_demanda"].between(
                pd.Timestamp(d_ini), pd.Timestamp(d_fim) + pd.Timedelta(hours=23, minutes=59, seconds=59)
            )
        )
    elif isinstance(periodo, date):
        mask &= (
            df_full["data_ultima_demanda"].isna()
            | (df_full["data_ultima_demanda"] >= pd.Timestamp(periodo))
            & (df_full["data_ultima_demanda"] < pd.Timestamp(periodo) + pd.Timedelta(days=1))
        )

    if status_sel:
        mask &= df_full["status_exito"].isin(status_sel)
    else:
        mask &= False

    if assunto_sel:
        mask &= df_full["assunto_label"].isin(assunto_sel)

    return df_full[mask].copy(), {
        "nucleo":     nucleo_sel,
        "procurador": procurador_sel,
        "periodo":    periodo,
    }


def _filtrar_demandas_frente2(df_dem_full: pd.DataFrame, filtros: dict) -> pd.DataFrame:
    """Filtra df_dem_full diretamente pelas seleções da sidebar, usando as
    datas nativas da própria demanda (conclusao / entrada) — independente dos
    processos resultantes da Frente 1."""
    if df_dem_full.empty:
        return df_dem_full.copy()

    mask = pd.Series(True, index=df_dem_full.index)

    nucleo_sel     = filtros.get("nucleo", [])
    procurador_sel = filtros.get("procurador", [])
    periodo        = filtros.get("periodo")

    if nucleo_sel:
        mask &= df_dem_full["nucleo"].isin(nucleo_sel)

    if procurador_sel:
        mask &= df_dem_full["procurador"].isin(procurador_sel)

    if isinstance(periodo, (list, tuple)) and len(periodo) == 2:
        d_ini, d_fim = periodo
        data_ref = df_dem_full["conclusao"].fillna(df_dem_full["entrada"])
        ts_fim   = pd.Timestamp(d_fim) + pd.Timedelta(hours=23, minutes=59, seconds=59)
        mask &= data_ref.isna() | data_ref.between(pd.Timestamp(d_ini), ts_fim)
    elif isinstance(periodo, date):
        data_ref = df_dem_full["conclusao"].fillna(df_dem_full["entrada"])
        mask &= (
            data_ref.isna()
            | (
                (data_ref >= pd.Timestamp(periodo))
                & (data_ref < pd.Timestamp(periodo) + pd.Timedelta(days=1))
            )
        )

    return df_dem_full[mask].copy()

# =============================================================================
# GRÁFICOS — TAB 1  (Panorama Geral)
# =============================================================================

def _donut(df: pd.DataFrame) -> None:
    agg = df.groupby("status_exito", as_index=False)["valor"].sum()

    fig = px.pie(
        agg, names="status_exito", values="valor",
        hole=0.68, color="status_exito", color_discrete_map=_STATUS_CORES,
    )
    fig.update_traces(
        textposition="outside",
        texttemplate="<b>%{percent:.1%}</b>",
        hovertemplate=(
            "<b>%{label}</b><br>Volume: R$ %{value:,.2f}"
            "<br>%{percent:.1%}<extra></extra>"
        ),
        marker=dict(line=dict(color="#FFFFFF", width=2)),
        textfont=dict(size=14, color="#111111"),
    )
    fig.update_layout(
        showlegend=True,
        legend=dict(
            orientation="h", y=-0.18, x=0.5, xanchor="center",
            font=dict(size=12, color="#212529"),
            bgcolor="rgba(0,0,0,0)",
        ),
        annotations=[dict(
            text=(
                f"<span style='font-size:26px;font-weight:900;color:#212529'>{len(df):,}</span><br>"
                f"<span style='font-size:11px;font-weight:600;letter-spacing:1px;color:#6C757D'>PROCESSOS</span>"
            ).replace(",", "."),
            x=0.5, y=0.5, showarrow=False,
        )],
        margin=dict(t=55, b=65, l=20, r=20),
        paper_bgcolor="#FFFFFF",
        plot_bgcolor="#FFFFFF",
        height=380,
    )
    st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False}, theme=None)


def _barras_top10(df: pd.DataFrame) -> None:
    agg = (
        df.groupby("assunto_label", as_index=False)["valor"]
        .sum().nlargest(10, "valor").sort_values("valor")
    )
    fig = px.bar(
        agg, x="valor", y="assunto_label", orientation="h", text="valor",
        color="valor",
        color_continuous_scale=["#FFCCCC", "#CC0000", "#5C0000"],
    )
    fig.update_traces(
        texttemplate="R$ %{x:,.0f}", textposition="outside", cliponaxis=False,
        hovertemplate="<b>%{y}</b><br>R$ %{x:,.2f}<extra></extra>",
        marker_line_width=0,
    )
    fig.update_layout(
        xaxis=dict(showgrid=False, showticklabels=False, zeroline=False, title="",
                   tickfont=dict(color="#212529")),
        yaxis=dict(showgrid=False, title="", tickfont=dict(size=10.5, color="#212529"),
                   automargin=True),
        font=dict(color="#212529"),
        coloraxis_showscale=False,
        plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
        margin=dict(t=10, b=10, l=5, r=100), height=360,
    )
    st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})

# =============================================================================
# GRÁFICOS — TAB 2  (Mapa e Comarcas)
# =============================================================================

def _barras_vitorias(df: pd.DataFrame) -> None:
    agg = (
        df[df["status_exito"] == "Vitória"]
        .groupby("comarca_limpa").size()
        .rename("vitorias").reset_index()
        .nlargest(15, "vitorias").sort_values("vitorias")
    )
    if agg.empty:
        st.info("Nenhuma vitória registrada na seleção atual.")
        return
    fig = px.bar(
        agg, x="vitorias", y="comarca_limpa", orientation="h", text="vitorias",
        color="vitorias",
        color_continuous_scale=["#C8F5C8", "#27AE60", "#145A32"],
    )
    fig.update_traces(
        textposition="outside", cliponaxis=False, marker_line_width=0,
        hovertemplate="<b>%{y}</b><br>Vitórias: %{x}<extra></extra>",
    )
    fig.update_layout(
        xaxis=dict(showgrid=False, showticklabels=False, zeroline=False, title="",
                   tickfont=dict(color="#212529")),
        yaxis=dict(showgrid=False, title="", tickfont=dict(size=10.5, color="#212529"),
                   automargin=True),
        font=dict(color="#212529"),
        coloraxis_showscale=False,
        plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
        margin=dict(t=10, b=10, l=5, r=45), height=470,
    )
    st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})


def _choropleth(df: pd.DataFrame, geojson: dict) -> None:
    vit = df[df["status_exito"] == "Vitória"].groupby("comarca_limpa").size().rename("vitorias")
    per = df[df["status_exito"] == "Perda"  ].groupby("comarca_limpa").size().rename("perdas")
    tot = df.groupby("comarca_limpa").size().rename("total")

    dm = pd.concat([vit, per, tot], axis=1).fillna(0).reset_index()
    dm["decididos"]    = dm["vitorias"] + dm["perdas"]
    dm["taxa_sucesso"] = (dm["vitorias"] / dm["decididos"] * 100).where(dm["decididos"] > 0)
    dm = dm[dm["comarca_limpa"].str.strip() != ""].dropna(subset=["taxa_sucesso"])

    if dm.empty:
        st.info("Não há processos decididos na seleção atual para gerar o mapa.")
        return

    fig = px.choropleth_mapbox(
        dm,
        geojson=geojson,
        locations="comarca_limpa",
        featureidkey="properties.nome_norm",
        color="taxa_sucesso",
        color_continuous_scale=_ESCALA_DIVERGENTE,
        range_color=[0, 100],
        mapbox_style="carto-positron",
        zoom=5.6,
        center={"lat": -22.0, "lon": -49.0},
        opacity=0.78,
        hover_name="comarca_limpa",
        custom_data=["vitorias", "perdas", "total", "taxa_sucesso"],
        labels={"taxa_sucesso": "Taxa Êxito (%)"},
    )
    fig.update_traces(
        hovertemplate=(
            "<b>%{hovertext}</b><br>"
            "Taxa de Êxito: <b>%{customdata[3]:.1f}%</b><br>"
            "Vitórias: %{customdata[0]:.0f}<br>"
            "Perdas: %{customdata[1]:.0f}<br>"
            "Total: %{customdata[2]:.0f}"
            "<extra></extra>"
        )
    )
    fig.update_layout(
        margin=dict(t=0, b=0, l=0, r=0),
        coloraxis_colorbar=dict(
            title=dict(text="Taxa<br>Êxito (%)", font=dict(color="#212529")),
            ticksuffix="%",
            len=0.65, thickness=14,
            tickvals=[0, 25, 50, 75, 100],
            tickfont=dict(color="#212529"),
        ),
        font=dict(color="#212529"),
        height=500,
        paper_bgcolor="#FFFFFF",
    )
    st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})


def _tabela_comarcas(df: pd.DataFrame) -> None:
    vit = df[df["status_exito"] == "Vitória"].groupby("comarca_limpa").size().rename("Vitórias")
    per = df[df["status_exito"] == "Perda"  ].groupby("comarca_limpa").size().rename("Perdas")
    tot = df.groupby("comarca_limpa").size().rename("Total")
    val = df.groupby("comarca_limpa")["valor"].sum().rename("Valor Total")

    dm = pd.concat([tot, vit, per, val], axis=1).fillna(0).reset_index()
    dm["Decididos"]      = dm["Vitórias"] + dm["Perdas"]
    dm["Taxa Êxito (%)"] = (
        dm["Vitórias"] / dm["Decididos"] * 100
    ).where(dm["Decididos"] > 0).round(1)
    dm["Valor Total"]    = dm["Valor Total"].apply(_brl)
    dm                   = dm.rename(columns={"comarca_limpa": "Comarca"})
    dm                   = dm.sort_values("Total", ascending=False)

    st.dataframe(
        dm[["Comarca", "Total", "Vitórias", "Perdas", "Decididos",
            "Taxa Êxito (%)", "Valor Total"]],
        use_container_width=True, hide_index=True, height=360,
        column_config={
            "Comarca":        st.column_config.TextColumn(width="medium"),
            "Total":          st.column_config.NumberColumn(format="%d", width="small"),
            "Vitórias":       st.column_config.NumberColumn(format="%d", width="small"),
            "Perdas":         st.column_config.NumberColumn(format="%d", width="small"),
            "Decididos":      st.column_config.NumberColumn(format="%d", width="small"),
            "Taxa Êxito (%)": st.column_config.NumberColumn(format="%.1f%%", width="small"),
            "Valor Total":    st.column_config.TextColumn(width="medium"),
        },
    )

# =============================================================================
# FRENTE 1 — ANÁLISE FINANCEIRA E RESULTADOS
# =============================================================================

def _subtab_f1_panorama(df: pd.DataFrame) -> None:
    if df.empty:
        st.warning("Nenhum processo corresponde à seleção. Ajuste os filtros.", icon="⚠️")
        return

    kpi = _kpis(df)

    c1, c2, c3, c4 = st.columns(4, gap="small")
    with c1:
        st.metric(
            "Total de Processos",
            f"{kpi['total']:,}".replace(",", "."),
            delta=f"{kpi['vitorias']}V · {kpi['perdas']}P · {kpi['decididos']} decididos",
            delta_color="off",
        )
    with c2:
        st.metric(
            "Valor em Risco (Em Andamento)",
            _brl(kpi["em_risco"]),
            delta=f"{(df['status_exito']=='Em Andamento').sum():,} em curso".replace(",", "."),
            delta_color="off",
        )
    with c3:
        st.metric(
            "Economia / Êxito Obtido",
            _brl(kpi["economia"]),
            delta=f"↑ {kpi['vitorias']} favoráveis ao Estado",
            delta_color="normal",
        )
    with c4:
        st.metric(
            "Taxa de Êxito (decididos)",
            f"{kpi['taxa']:.1f}%",
            delta=f"{kpi['vitorias']}V / {kpi['perdas']}P de {kpi['decididos']}",
            delta_color="normal" if kpi["taxa"] >= 50 else "inverse",
        )

    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    col_e, col_d = st.columns([1, 2], gap="medium")
    with col_e:
        st.markdown('<p class="sec-title">Volume Financeiro por Status</p>', unsafe_allow_html=True)
        _donut(df)
    with col_d:
        st.markdown('<p class="sec-title">Top 10 Assuntos por Volume Financeiro</p>', unsafe_allow_html=True)
        _barras_top10(df)

    nucleos_sel = st.session_state.get("f_nucleo", [])
    if nucleos_sel:
        titulo_fin   = "Desempenho Financeiro por Procurador"
        subtitulo_fin = "Valor acumulado por procurador segregado por desfecho processual"
        group_col, label_col = "procurador", "Procurador"
    else:
        titulo_fin   = "Desempenho Financeiro por Núcleo"
        subtitulo_fin = "Valor acumulado por núcleo segregado por desfecho processual"
        group_col, label_col = "nucleo", "Núcleo"

    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    st.markdown(f'<p class="sec-title">{titulo_fin}</p>', unsafe_allow_html=True)
    st.markdown(f'<p class="sec-sub">{subtitulo_fin}</p>', unsafe_allow_html=True)

    _tabela_financeira_grupo(df, group_col, label_col)


def _tabela_financeira_grupo(df: pd.DataFrame, group_col: str, label_col: str) -> None:
    risco = (
        df[df["status_exito"] == "Em Andamento"]
        .groupby(group_col)["valor"].sum()
        .rename("_risco")
    )
    salvo = (
        df[df["status_exito"] == "Vitória"]
        .groupby(group_col)["valor"].sum()
        .rename("_salvo")
    )
    perdido = (
        df[df["status_exito"] == "Perda"]
        .groupby(group_col)["valor"].sum()
        .rename("_perdido")
    )

    fp = pd.concat([risco, salvo, perdido], axis=1).fillna(0.0).reset_index()
    fp.columns = [label_col, "_risco", "_salvo", "_perdido"]

    placeholder = {"procurador": "(sem procurador)", "nucleo": "(sem núcleo)"}.get(group_col)
    if placeholder:
        fp = fp[fp[label_col] != placeholder]

    decidido_val = fp["_salvo"] + fp["_perdido"]
    fp["Taxa Êxito Fin. (%)"] = (
        (fp["_salvo"] / decidido_val * 100)
        .where(decidido_val > 0)
        .fillna(0.0)
        .round(1)
    )
    fp = fp.sort_values("_risco", ascending=False)

    fp["Valor em Risco"]    = fp["_risco"].apply(_brl)
    fp["Valor Salvo"]       = fp["_salvo"].apply(_brl)
    fp["Valor Perdido"]     = fp["_perdido"].apply(_brl)

    fp_show = fp[[label_col, "Valor em Risco", "Valor Salvo", "Valor Perdido", "Taxa Êxito Fin. (%)"]].copy()

    st.dataframe(
        fp_show,
        use_container_width=True,
        hide_index=True,
        height=380,
        column_config={
            label_col:             st.column_config.TextColumn(width="large"),
            "Valor em Risco":      st.column_config.TextColumn(width="medium"),
            "Valor Salvo":         st.column_config.TextColumn(width="medium"),
            "Valor Perdido":       st.column_config.TextColumn(width="medium"),
            "Taxa Êxito Fin. (%)": st.column_config.ProgressColumn(
                label="Taxa Êxito Fin. (%)",
                format="%.1f%%",
                min_value=0.0,
                max_value=100.0,
                width="medium",
            ),
        },
    )
    rotulo_plural = "procuradores" if group_col == "procurador" else "núcleos"
    st.markdown(
        f"<small style='color:#ADB5BD'>{len(fp_show):,} {rotulo_plural}</small>".replace(",", "."),
        unsafe_allow_html=True,
    )


def _subtab_f1_exito_tese(df: pd.DataFrame) -> None:
    if df.empty:
        st.warning("Nenhum processo corresponde à seleção. Ajuste os filtros.", icon="⚠️")
        return

    vit = df[df["status_exito"] == "Vitória"].groupby("assunto_label").size().rename("Vitórias")
    per = df[df["status_exito"] == "Perda"  ].groupby("assunto_label").size().rename("Perdas")
    tot = df.groupby("assunto_label").size().rename("Total")
    val = df.groupby("assunto_label")["valor"].sum().rename("Valor Total")

    dm = pd.concat([tot, vit, per, val], axis=1).fillna(0).reset_index()
    dm["Decididos"]      = dm["Vitórias"] + dm["Perdas"]
    dm["Taxa Êxito (%)"] = (dm["Vitórias"] / dm["Decididos"] * 100).where(dm["Decididos"] > 0).round(1)
    dm = dm[dm["Decididos"] > 0].copy()

    if dm.empty:
        st.info("Nenhuma tese possui processos decididos na seleção atual.", icon="ℹ️")
        return

    kpi1, kpi2, kpi3 = st.columns(3, gap="small")
    with kpi1:
        st.metric("Teses Analisadas", f"{len(dm):,}".replace(",", "."))
    with kpi2:
        melhor = dm.nlargest(1, "Taxa Êxito (%)").iloc[0]
        st.metric("Melhor Tese (Taxa)", f"{melhor['Taxa Êxito (%)']:.1f}%",
                  delta=str(melhor["assunto_label"])[:40], delta_color="off")
    with kpi3:
        pior = dm.nsmallest(1, "Taxa Êxito (%)").iloc[0]
        st.metric("Tese de Maior Risco", f"{pior['Taxa Êxito (%)']:.1f}%",
                  delta=str(pior["assunto_label"])[:40], delta_color="off")

    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    dm_chart = dm.nlargest(15, "Decididos").sort_values("Taxa Êxito (%)")
    fig = px.bar(
        dm_chart, x="Taxa Êxito (%)", y="assunto_label", orientation="h",
        text="Taxa Êxito (%)", color="Taxa Êxito (%)",
        color_continuous_scale=_ESCALA_DIVERGENTE,
        range_color=[0, 100],
        custom_data=["Vitórias", "Perdas", "Decididos", "Total"],
    )
    fig.update_traces(
        texttemplate="%{x:.1f}%", textposition="outside", cliponaxis=False,
        marker_line_width=0,
        hovertemplate=(
            "<b>%{y}</b><br>Taxa Êxito: %{x:.1f}%<br>"
            "Vitórias: %{customdata[0]:.0f} | Perdas: %{customdata[1]:.0f}<br>"
            "Decididos: %{customdata[2]:.0f} / Total: %{customdata[3]:.0f}"
            "<extra></extra>"
        ),
    )
    fig.update_layout(
        xaxis=dict(showgrid=False, showticklabels=False, zeroline=False, title="", range=[0, 120]),
        yaxis=dict(showgrid=False, title="", tickfont=dict(size=10.5, color="#212529"), automargin=True),
        coloraxis_showscale=False,
        plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
        margin=dict(t=10, b=10, l=5, r=70), height=420,
        font=dict(color="#212529"),
    )
    st.markdown('<p class="sec-title">Taxa de Êxito por Tese (Top 15 por Volume Decidido)</p>',
                unsafe_allow_html=True)
    st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False}, theme=None)

    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    st.markdown('<p class="sec-title">Tabela Completa por Tese</p>', unsafe_allow_html=True)

    dm_tab = dm.copy()
    dm_tab["Valor Total"] = dm_tab["Valor Total"].apply(_brl)
    dm_tab = dm_tab.rename(columns={"assunto_label": "Tese / Assunto"}).sort_values(
        "Decididos", ascending=False
    )
    st.dataframe(
        dm_tab[["Tese / Assunto", "Total", "Vitórias", "Perdas", "Decididos", "Taxa Êxito (%)", "Valor Total"]],
        use_container_width=True, hide_index=True, height=380,
        column_config={
            "Tese / Assunto":  st.column_config.TextColumn(width="large"),
            "Total":           st.column_config.NumberColumn(format="%d", width="small"),
            "Vitórias":        st.column_config.NumberColumn(format="%d", width="small"),
            "Perdas":          st.column_config.NumberColumn(format="%d", width="small"),
            "Decididos":       st.column_config.NumberColumn(format="%d", width="small"),
            "Taxa Êxito (%)":  st.column_config.NumberColumn(format="%.1f%%", width="small"),
            "Valor Total":     st.column_config.TextColumn(width="medium"),
        },
    )


def _subtab_f1_detalhamento(df: pd.DataFrame) -> None:
    if df.empty:
        st.warning("Nenhum processo corresponde à seleção. Ajuste os filtros.", icon="⚠️")
        return

    st.markdown('<p class="sec-title">Detalhamento de Processos</p>', unsafe_allow_html=True)

    cols_disp = [c for c in _COLUNAS_TABELA if c in df.columns]
    df_tab = df[cols_disp].rename(columns=_COLUNAS_TABELA).copy()
    df_tab["Valor (R$)"] = df_tab["Valor (R$)"].apply(_brl)
    if "Data" in df_tab.columns:
        df_tab["Data"] = (
            pd.to_datetime(df_tab["Data"], errors="coerce")
            .dt.strftime("%d/%m/%Y")
            .fillna("")
        )

    st.dataframe(
        df_tab, use_container_width=True, hide_index=True, height=540,
        column_config={
            "Nº do Processo":  st.column_config.TextColumn(width="medium"),
            "Valor (R$)":      st.column_config.TextColumn(width="small"),
            "Assunto":         st.column_config.TextColumn(width="large"),
            "Procurador":      st.column_config.TextColumn(width="medium"),
            "Vara":            st.column_config.TextColumn(width="medium"),
            "Última Demanda":  st.column_config.TextColumn(width="large"),
            "Data":            st.column_config.TextColumn(width="small"),
            "Status":          st.column_config.TextColumn(width="small"),
        },
    )
    st.markdown(
        f"<small style='color:#ADB5BD'>{len(df_tab):,} registros</small>".replace(",", "."),
        unsafe_allow_html=True,
    )


def _subtab_f1_linha_tempo(df: pd.DataFrame) -> None:
    if df.empty:
        st.warning("Nenhum processo corresponde à seleção. Ajuste os filtros.", icon="⚠️")
        return

    base = (
        df[df["status_exito"].isin(["Vitória", "Perda"])]
        .dropna(subset=["data_ultima_demanda"])
        .copy()
    )
    if base.empty:
        st.info("Nenhuma vitória ou perda com data registrada na seleção atual.", icon="ℹ️")
        return

    base["mes"]       = base["data_ultima_demanda"].dt.to_period("M").dt.to_timestamp()
    base["mes_label"] = base["mes"].dt.strftime("%m/%Y")

    agg = (
        base.groupby(["mes", "mes_label", "status_exito"], as_index=False)["valor"]
        .sum()
        .sort_values("mes")
    )
    # Perda fica negativa para desenhar abaixo da linha zero; Vitória acima —
    # cria a linha do tempo centralizada solicitada pelo usuário.
    agg["valor_eixo"] = agg.apply(
        lambda r: -r["valor"] if r["status_exito"] == "Perda" else r["valor"], axis=1
    )

    ordem_meses  = agg.sort_values("mes")["mes_label"].drop_duplicates().tolist()
    label_p_mes  = dict(zip(agg["mes_label"], agg["mes"]))
    n_meses      = len(ordem_meses)

    st.markdown('<p class="sec-title">Linha do Tempo — Vitórias e Perdas por Mês</p>',
                unsafe_allow_html=True)

    fig = px.bar(
        agg, x="mes_label", y="valor_eixo", color="status_exito",
        barmode="relative",
        color_discrete_map=_STATUS_CORES,
        custom_data=["valor"],
        category_orders={"mes_label": ordem_meses},
        labels={"mes_label": "", "valor_eixo": "Valor (R$)", "status_exito": "Status"},
    )
    fig.update_traces(
        marker_line_width=0,
        hovertemplate="<b>%{x}</b><br>%{customdata[0]:,.2f}<extra></extra>",
    )
    fig.update_xaxes(
        type="category", showgrid=False, automargin=True,
        tickangle=(-40 if n_meses > 8 else 0),
        tickfont=dict(size=10.5, color="#212529"),
    )
    fig.update_layout(
        bargap=0.35, bargroupgap=0.08,
        yaxis=dict(
            showgrid=True, gridcolor="#EEEEEE", title="Valor (R$)",
            zeroline=True, zerolinewidth=2, zerolinecolor="#212529",
            tickprefix="R$ ", tickformat="~s", automargin=True,
        ),
        plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
        margin=dict(t=40, b=60, l=10, r=10), height=380,
        font=dict(color="#212529"),
        legend=dict(orientation="h", yanchor="bottom", y=1.02, x=0),
        clickmode="event+select",
    )

    evento = st.plotly_chart(
        fig, use_container_width=True, theme=None,
        key="f1_linha_tempo_chart", on_select="rerun",
        selection_mode=("points", "box", "lasso"),
        config={"displayModeBar": False},
    )

    pontos = (evento or {}).get("selection", {}).get("points", [])
    meses_sel = sorted({label_p_mes[p["x"]] for p in pontos if p.get("x") in label_p_mes})

    if meses_sel:
        meses_dt = pd.to_datetime(meses_sel)
        detalhe  = base[base["mes"].isin(meses_dt)]
        st.caption(
            f"🔎 Exibindo processos de {meses_dt.min().strftime('%m/%Y')} a "
            f"{meses_dt.max().strftime('%m/%Y')} — clique nas barras novamente "
            f"para ajustar a seleção, ou em uma área vazia do gráfico para limpá-la."
        )
    else:
        detalhe = base
        st.caption(
            f"Exibindo todos os meses disponíveis "
            f"({base['mes'].min().strftime('%m/%Y')} a {base['mes'].max().strftime('%m/%Y')}). "
            f"Clique em uma ou mais barras do gráfico para filtrar por mês."
        )

    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    st.markdown('<p class="sec-title">Processos do Período Selecionado</p>', unsafe_allow_html=True)

    cols_disp = [c for c in _COLUNAS_TABELA if c in detalhe.columns]
    df_tab = detalhe[cols_disp].rename(columns=_COLUNAS_TABELA).copy()
    df_tab["Valor (R$)"] = df_tab["Valor (R$)"].apply(_brl)
    if "Data" in df_tab.columns:
        df_tab["Data"] = (
            pd.to_datetime(df_tab["Data"], errors="coerce")
            .dt.strftime("%d/%m/%Y")
            .fillna("")
        )

    st.dataframe(
        df_tab, use_container_width=True, hide_index=True, height=420,
        column_config={
            "Nº do Processo":  st.column_config.TextColumn(width="medium"),
            "Valor (R$)":      st.column_config.TextColumn(width="small"),
            "Assunto":         st.column_config.TextColumn(width="large"),
            "Procurador":      st.column_config.TextColumn(width="medium"),
            "Vara":            st.column_config.TextColumn(width="medium"),
            "Última Demanda":  st.column_config.TextColumn(width="large"),
            "Data":            st.column_config.TextColumn(width="small"),
            "Status":          st.column_config.TextColumn(width="small"),
        },
    )
    st.markdown(
        f"<small style='color:#ADB5BD'>{len(df_tab):,} registros</small>".replace(",", "."),
        unsafe_allow_html=True,
    )


_ESTAGIOS_ORDEM = [
    "Transitado em Julgado",
    "Fase Recursal (Tribunais)",
    "Sentenciado (1ª Instância)",
    "Fase Inicial / Conhecimento",
]
_ESTAGIOS_CORES = {
    "Transitado em Julgado":         "#27AE60",
    "Fase Recursal (Tribunais)":     "#0070CC",
    "Sentenciado (1ª Instância)":    "#E07B00",
    "Fase Inicial / Conhecimento":   "#ADB5BD",
}


def _classif_estagio(demandas_iter) -> str:
    lista = [str(d) for d in demandas_iter if d and not pd.isna(d)]
    if any("Trânsito em julgado" in d for d in lista):
        return "Transitado em Julgado"
    if any(kw in d for d in lista for kw in ("Acórdão", "contrarrazões", "contraminuta")):
        return "Fase Recursal (Tribunais)"
    if any("Intimação de Sentença" in d for d in lista):
        return "Sentenciado (1ª Instância)"
    return "Fase Inicial / Conhecimento"


def _subtab_f1_estagio(df: pd.DataFrame, df_dem_full: pd.DataFrame) -> None:
    if df.empty:
        st.warning("Nenhum processo corresponde à seleção. Ajuste os filtros.", icon="⚠️")
        return

    if df_dem_full.empty or "demanda" not in df_dem_full.columns:
        st.info(
            "Dados de demandas não disponíveis para classificar o estágio processual. "
            "Faça o upload do Relatório de Demandas.",
            icon="ℹ️",
        )
        return

    # Classificar cada processo_base pelo seu estágio mais avançado
    estagio_map = (
        df_dem_full.groupby("processo_base")["demanda"]
        .apply(_classif_estagio)
        .reset_index(name="Estágio")
    )

    # Cruzar com processos filtrados (df usa "processo" como PK)
    df_est = df.merge(
        estagio_map,
        left_on="processo",
        right_on="processo_base",
        how="left",
    )
    df_est["Estágio"] = df_est["Estágio"].fillna("Fase Inicial / Conhecimento")

    # ── KPIs por estágio ───────────────────────────────────────────────────
    agg_est = (
        df_est.groupby("Estágio")["valor"]
        .sum()
        .reindex(_ESTAGIOS_ORDEM, fill_value=0.0)
    )

    c1, c2, c3, c4 = st.columns(4, gap="small")
    for col, estagio in zip([c1, c2, c3, c4], _ESTAGIOS_ORDEM):
        n_proc = int((df_est["Estágio"] == estagio).sum())
        with col:
            st.metric(
                estagio,
                _brl(agg_est.get(estagio, 0.0)),
                delta=f"{n_proc:,} processos".replace(",", "."),
                delta_color="off",
            )

    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # ── Gráfico de barras verticais ────────────────────────────────────────
    st.markdown(
        '<p class="sec-title">Volume Financeiro por Estágio Processual</p>',
        unsafe_allow_html=True,
    )

    agg_chart = (
        df_est.groupby("Estágio")
        .agg(valor_total=("valor", "sum"), processos=("processo", "nunique"))
        .reindex(_ESTAGIOS_ORDEM, fill_value=0)
        .reset_index()
    )
    agg_chart["Cor"] = agg_chart["Estágio"].map(_ESTAGIOS_CORES)

    fig = px.bar(
        agg_chart,
        x="Estágio",
        y="valor_total",
        color="Estágio",
        color_discrete_map=_ESTAGIOS_CORES,
        text="valor_total",
        custom_data=["processos"],
        category_orders={"Estágio": _ESTAGIOS_ORDEM},
    )
    fig.update_traces(
        texttemplate="%{y:,.0f}",
        textposition="outside",
        textfont=dict(size=11),
        cliponaxis=False,
        marker_line_width=0,
        hovertemplate=(
            "<b>%{x}</b><br>"
            "Volume: R$ %{y:,.2f}<br>"
            "Processos: %{customdata[0]}"
            "<extra></extra>"
        ),
    )

    max_val = float(agg_chart["valor_total"].max()) if not agg_chart.empty else 0.0
    y_max   = max_val * 1.25 if max_val > 0 else 1.0

    fig.update_layout(
        xaxis=dict(title="", tickfont=dict(size=10.5, color="#212529"), showgrid=False,
                   automargin=True),
        yaxis=dict(title="Valor (R$)", tickfont=dict(color="#212529"),
                   showgrid=True, gridcolor="#F0F0F0", showticklabels=False,
                   range=[0, y_max], automargin=True),
        showlegend=False,
        plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
        margin=dict(t=50, b=90, l=10, r=10), height=440,
        font=dict(color="#212529"),
    )
    st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False}, theme=None)

    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    # ── Filtro e tabela por estágio ────────────────────────────────────────
    st.markdown('<p class="sec-title">Processos por Estágio</p>', unsafe_allow_html=True)

    estagios_disponiveis = [e for e in _ESTAGIOS_ORDEM if e in df_est["Estágio"].values]
    estagio_sel = st.selectbox(
        "Filtrar estágio",
        options=estagios_disponiveis,
        key="f1_estagio_sel",
        label_visibility="collapsed",
    )

    df_estagio_tab = df_est[df_est["Estágio"] == estagio_sel].copy()

    cols_show = [c for c in ["processo", "procurador", "assunto_label", "ult_demanda", "valor"]
                 if c in df_estagio_tab.columns]
    df_tab = df_estagio_tab[cols_show].rename(columns={
        "processo":      "Nº do Processo",
        "procurador":    "Procurador",
        "assunto_label": "Assunto",
        "ult_demanda":   "Última Demanda",
        "valor":         "Valor (R$)",
    }).copy()
    df_tab["Valor (R$)"] = df_tab["Valor (R$)"].apply(_brl)
    df_tab = df_tab.sort_values("Nº do Processo")

    st.dataframe(
        df_tab,
        use_container_width=True,
        hide_index=True,
        height=420,
        column_config={
            "Nº do Processo": st.column_config.TextColumn(width="medium"),
            "Procurador":     st.column_config.TextColumn(width="medium"),
            "Assunto":        st.column_config.TextColumn(width="large"),
            "Última Demanda": st.column_config.TextColumn(width="large"),
            "Valor (R$)":     st.column_config.TextColumn(width="small"),
        },
    )
    st.markdown(
        f"<small style='color:#ADB5BD'>{len(df_tab):,} processos em «{estagio_sel}»</small>".replace(",", "."),
        unsafe_allow_html=True,
    )


def _subtab_f1_mapa(df: pd.DataFrame) -> None:
    if df.empty:
        st.warning("Nenhum processo corresponde à seleção. Ajuste os filtros.", icon="⚠️")
        return

    col_bar, col_map = st.columns([1, 2], gap="medium")

    with col_bar:
        st.markdown('<p class="sec-title">Top Comarcas · Vitórias</p>', unsafe_allow_html=True)
        st.markdown(
            '<p class="sec-sub">Volume absoluto de resultados favoráveis ao Estado</p>',
            unsafe_allow_html=True,
        )
        _barras_vitorias(df)

    with col_map:
        st.markdown('<p class="sec-title">Mapa de Calor · Taxa de Êxito por Comarca</p>', unsafe_allow_html=True)
        st.markdown(
            '<p class="sec-sub">'
            'Vitórias ÷ (Vitórias + Perdas) × 100 &nbsp;·&nbsp; '
            '<span style="color:#CC0000">■</span> 0% &nbsp;'
            '<span style="color:#F5C518">■</span> 50% &nbsp;'
            '<span style="color:#27AE60">■</span> 100%'
            '</p>',
            unsafe_allow_html=True,
        )
        geojson = _carregar_geojson()
        if geojson is None:
            st.warning(
                "**Mapa indisponível** — não foi possível baixar o GeoJSON dos municípios de SP. "
                "Verifique sua conexão e recarregue a página.",
                icon="🌐",
            )
        else:
            _choropleth(df, geojson)

    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    st.markdown('<p class="sec-title">Análise Detalhada por Comarca</p>', unsafe_allow_html=True)
    _tabela_comarcas(df)


# =============================================================================
# FRENTE 2 — GESTÃO OPERACIONAL E PRODUTIVIDADE
# =============================================================================

def _subtab_f2_nucleos(df: pd.DataFrame, df_dem: pd.DataFrame) -> None:
    if df_dem.empty:
        st.info(
            "Nenhum dado de demandas disponível para a seleção atual. "
            "Faça o upload do Relatório de Demandas ou ajuste os filtros.",
            icon="ℹ️",
        )
        return

    c1, c2, c3, c4 = st.columns(4, gap="small")
    with c1:
        st.metric(
            "Núcleos Ativos",
            f"{df_dem['nucleo'].nunique():,}".replace(",", "."),
        )
    with c2:
        st.metric(
            "Total de Demandas",
            f"{len(df_dem):,}".replace(",", "."),
        )
    with c3:
        st.metric(
            "Total de Horas",
            f"{df_dem['horas'].sum():,.1f}".replace(",", "."),
        )
    with c4:
        media_h = df_dem["horas"].mean() if len(df_dem) > 0 else 0.0
        st.metric("Média Horas/Demanda", f"{media_h:.1f}")

    st.markdown('<hr class="divider">', unsafe_allow_html=True)

    col_nuc, col_tipo = st.columns(2, gap="medium")

    with col_nuc:
        st.markdown(
            '<p class="sec-title">Ranking de Núcleos por Nº de Demandas</p>',
            unsafe_allow_html=True,
        )
        agg_nuc = (
            df_dem.groupby("nucleo", as_index=False)
            .size()
            .rename(columns={"size": "demandas"})
            .sort_values("demandas", ascending=False)
        )
        fig = px.pie(
            agg_nuc, names="nucleo", values="demandas", hole=0.45,
        )
        fig.update_traces(
            textposition="inside", textinfo="percent",
            hovertemplate="<b>%{label}</b><br>Demandas: %{value}<br>%{percent}<extra></extra>",
            marker=dict(line=dict(color="#FFFFFF", width=1.5)),
        )
        fig.update_layout(
            showlegend=True,
            legend=dict(orientation="h", y=-0.2, x=0.5, xanchor="center",
                        font=dict(size=9.5, color="#212529")),
            margin=dict(t=10, b=10, l=10, r=10), height=430,
            paper_bgcolor="#FFFFFF",
        )
        st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False}, theme=None)

    with col_tipo:
        st.markdown(
            '<p class="sec-title">Tipos de Demanda mais Frequentes (Top 15)</p>',
            unsafe_allow_html=True,
        )
        agg_tipo = (
            df_dem.groupby("demanda", as_index=False)
            .size()
            .rename(columns={"size": "qtd"})
            .nlargest(15, "qtd")
            .sort_values("qtd")
        )
        fig2 = px.bar(
            agg_tipo, x="qtd", y="demanda", orientation="h", text="qtd",
            color="qtd",
            color_continuous_scale=["#CCE5FF", "#0070CC", "#003D7A"],
        )
        fig2.update_traces(
            textposition="outside", cliponaxis=False, marker_line_width=0,
            hovertemplate="<b>%{y}</b><br>Qtd: %{x}<extra></extra>",
        )
        fig2.update_layout(
            xaxis=dict(showgrid=False, showticklabels=False, zeroline=False, title=""),
            yaxis=dict(showgrid=False, title="", tickfont=dict(size=9.5, color="#212529"),
                       automargin=True),
            coloraxis_showscale=False,
            plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
            margin=dict(t=10, b=10, l=5, r=45), height=430,
            font=dict(color="#212529"),
        )
        st.plotly_chart(fig2, use_container_width=True, config={"displayModeBar": False})

    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    st.markdown('<p class="sec-title">Detalhamento por Núcleo</p>', unsafe_allow_html=True)

    agg_det = (
        df_dem.groupby("nucleo")
        .agg(
            Demandas=("demanda", "count"),
            Horas_Total=("horas", "sum"),
            Horas_Medias=("horas", "mean"),
            Processos_Unicos=("processo_base", "nunique"),
            Procuradores_Unicos=("procurador", "nunique"),
        )
        .reset_index()
        .rename(columns={
            "nucleo":              "Núcleo",
            "Horas_Total":         "Horas Total",
            "Horas_Medias":        "Horas Médias",
            "Processos_Unicos":    "Processos Únicos",
            "Procuradores_Unicos": "Procuradores",
        })
        .sort_values("Demandas", ascending=False)
    )
    agg_det["Horas Total"]  = agg_det["Horas Total"].round(1)
    agg_det["Horas Médias"] = agg_det["Horas Médias"].round(1)

    if not df.empty and "nucleo" in df.columns:
        risco = (
            df[df["status_exito"] == "Em Andamento"]
            .groupby("nucleo")
            .agg(
                Processos_Ativos=("processo", "nunique"),
                Valor_Conducao=("valor", "sum"),
            )
            .reset_index()
            .rename(columns={"nucleo": "Núcleo"})
        )
        agg_det = agg_det.merge(risco, on="Núcleo", how="left")
        agg_det["Processos_Ativos"] = agg_det["Processos_Ativos"].fillna(0).astype(int)
        agg_det["Valor Sob Condução"] = agg_det["Valor_Conducao"].fillna(0.0).apply(_brl)
        agg_det = agg_det.drop(columns=["Valor_Conducao"])
        cols_show = ["Núcleo", "Procuradores", "Demandas", "Horas Total", "Horas Médias",
                     "Processos Únicos", "Processos_Ativos", "Valor Sob Condução"]
        col_cfg = {
            "Núcleo":               st.column_config.TextColumn(width="large"),
            "Procuradores":         st.column_config.NumberColumn(format="%d",   width="small"),
            "Demandas":             st.column_config.NumberColumn(format="%d",   width="small"),
            "Horas Total":          st.column_config.NumberColumn(format="%.1f", width="small"),
            "Horas Médias":         st.column_config.NumberColumn(format="%.1f", width="small"),
            "Processos Únicos":     st.column_config.NumberColumn(format="%d",   width="small"),
            "Processos_Ativos":     st.column_config.NumberColumn(format="%d",   width="small",
                                                                  label="Processos Ativos"),
            "Valor Sob Condução":   st.column_config.TextColumn(width="medium"),
        }
    else:
        cols_show = ["Núcleo", "Procuradores", "Demandas", "Horas Total", "Horas Médias", "Processos Únicos"]
        col_cfg = {
            "Núcleo":           st.column_config.TextColumn(width="large"),
            "Procuradores":     st.column_config.NumberColumn(format="%d",   width="small"),
            "Demandas":         st.column_config.NumberColumn(format="%d",   width="small"),
            "Horas Total":      st.column_config.NumberColumn(format="%.1f", width="small"),
            "Horas Médias":     st.column_config.NumberColumn(format="%.1f", width="small"),
            "Processos Únicos": st.column_config.NumberColumn(format="%d",   width="small"),
        }

    st.dataframe(
        agg_det[cols_show], use_container_width=True, hide_index=True, height=420,
        column_config=col_cfg,
    )


def _subtab_f2_fluxo(df_dem: pd.DataFrame) -> None:
    if df_dem.empty:
        st.info("Nenhum dado de demandas disponível.", icon="ℹ️")
        return

    cutoff = pd.Timestamp("2026-01-01")
    ent = df_dem.dropna(subset=["entrada"])
    con = df_dem.dropna(subset=["conclusao"])

    ent_2026 = ent[ent["entrada"] >= cutoff].copy()
    con_2026 = con[con["conclusao"] >= cutoff].copy()

    if ent_2026.empty and con_2026.empty:
        st.info("Nenhum registro de entrada ou conclusão a partir de 01/01/2026.", icon="ℹ️")
        return

    dias_ent = ent_2026.groupby(ent_2026["entrada"].dt.normalize()).size().rename("Entradas")
    dias_con = con_2026.groupby(con_2026["conclusao"].dt.normalize()).size().rename("Conclusões")

    dm = pd.concat([dias_ent, dias_con], axis=1).fillna(0).reset_index()
    dm = dm.rename(columns={"index": "Data"})
    if "Data" not in dm.columns:
        dm = dm.rename(columns={dm.columns[0]: "Data"})
    dm = dm.sort_values("Data")
    dm["Saldo"]          = dm["Conclusões"] - dm["Entradas"]
    dm["Saldo Acumulado"] = dm["Saldo"].cumsum()

    kc1, kc2, kc3 = st.columns(3, gap="small")
    with kc1:
        st.metric("Total Entradas (2026)", f"{int(dm['Entradas'].sum()):,}".replace(",", "."))
    with kc2:
        st.metric("Total Conclusões (2026)", f"{int(dm['Conclusões'].sum()):,}".replace(",", "."))
    with kc3:
        saldo_final = int(dm["Saldo Acumulado"].iloc[-1]) if len(dm) > 0 else 0
        st.metric("Saldo Acumulado", f"{saldo_final:+,}".replace(",", "."),
                  delta_color="normal" if saldo_final >= 0 else "inverse")

    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    st.markdown('<p class="sec-title">Entradas vs Conclusões Diárias (a partir de 01/01/2026)</p>',
                unsafe_allow_html=True)

    fig = px.line(
        dm, x="Data", y=["Entradas", "Conclusões"],
        color_discrete_map={"Entradas": "#CC0000", "Conclusões": "#27AE60"},
    )
    fig.update_traces(mode="lines+markers", marker=dict(size=4))
    fig.update_layout(
        xaxis=dict(title="", tickfont=dict(color="#212529"), showgrid=False),
        yaxis=dict(title="Demandas", tickfont=dict(color="#212529"), showgrid=True,
                   gridcolor="#F0F0F0"),
        legend=dict(title="", orientation="h", y=1.05, x=0, font=dict(color="#212529")),
        plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
        margin=dict(t=20, b=10, l=10, r=10), height=320,
        font=dict(color="#212529"),
    )
    st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False}, theme=None)

    st.markdown('<p class="sec-title">Saldo Acumulado (Conclusões − Entradas)</p>',
                unsafe_allow_html=True)
    fig2 = px.area(
        dm, x="Data", y="Saldo Acumulado",
        color_discrete_sequence=["#0070CC"],
    )
    fig2.add_hline(y=0, line_dash="dash", line_color="#ADB5BD")
    fig2.update_layout(
        xaxis=dict(title="", tickfont=dict(color="#212529"), showgrid=False),
        yaxis=dict(title="Saldo", tickfont=dict(color="#212529"), showgrid=True,
                   gridcolor="#F0F0F0"),
        plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
        margin=dict(t=10, b=10, l=10, r=10), height=280,
        font=dict(color="#212529"),
    )
    st.plotly_chart(fig2, use_container_width=True, config={"displayModeBar": False}, theme=None)

    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    st.markdown('<p class="sec-title">Tabela de Fluxo Diário</p>', unsafe_allow_html=True)

    tab_fluxo = dm[["Data", "Conclusões", "Entradas", "Saldo"]].copy()
    tab_fluxo = tab_fluxo.sort_values("Data", ascending=False)
    tab_fluxo["Data"] = pd.to_datetime(tab_fluxo["Data"]).dt.strftime("%d/%m/%Y")
    tab_fluxo = tab_fluxo.rename(columns={
        "Conclusões": "Total de Conclusões",
        "Entradas":   "Total de Entradas",
    })
    tab_fluxo["Total de Conclusões"] = tab_fluxo["Total de Conclusões"].astype(int)
    tab_fluxo["Total de Entradas"]   = tab_fluxo["Total de Entradas"].astype(int)
    tab_fluxo["Saldo"]               = tab_fluxo["Saldo"].astype(int)

    st.dataframe(
        tab_fluxo[["Data", "Total de Conclusões", "Total de Entradas", "Saldo"]],
        use_container_width=True, hide_index=True, height=400,
        column_config={
            "Data":                 st.column_config.TextColumn(width="small"),
            "Total de Conclusões":  st.column_config.NumberColumn(format="%d", width="small"),
            "Total de Entradas":    st.column_config.NumberColumn(format="%d", width="small"),
            "Saldo":                st.column_config.NumberColumn(format="%+d", width="small"),
        },
    )


def _subtab_f2_sazonalidade(df_dem: pd.DataFrame) -> None:
    if df_dem.empty:
        st.info("Nenhum dado de demandas disponível.", icon="ℹ️")
        return

    _dias_pt = ["Segunda", "Terça", "Quarta", "Quinta", "Sexta", "Sábado", "Domingo"]

    ent_validas = df_dem.dropna(subset=["entrada"])
    con_validas = df_dem.dropna(subset=["conclusao"])

    if ent_validas.empty and con_validas.empty:
        st.info("Sem datas de entrada ou conclusão para análise de sazonalidade.", icon="ℹ️")
        return

    kc1, kc2 = st.columns(2, gap="medium")

    with kc1:
        st.markdown('<p class="sec-title">Distribuição de Entradas por Dia da Semana</p>',
                    unsafe_allow_html=True)
        if not ent_validas.empty:
            ent_dow = ent_validas["entrada"].dt.dayofweek.value_counts().sort_index().reset_index()
            ent_dow.columns = ["dow", "qtd"]
            ent_dow["Dia"] = ent_dow["dow"].map(lambda x: _dias_pt[x])
            fig = px.bar(
                ent_dow, x="Dia", y="qtd", text="qtd",
                color="qtd",
                color_continuous_scale=["#FFCCCC", "#CC0000", "#5C0000"],
            )
            fig.update_traces(
                textposition="outside", cliponaxis=False, marker_line_width=0,
                hovertemplate="<b>%{x}</b><br>Entradas: %{y}<extra></extra>",
            )
            fig.update_layout(
                xaxis=dict(title="", tickfont=dict(color="#212529"), showgrid=False,
                           categoryorder="array", categoryarray=_dias_pt),
                yaxis=dict(title="Qtd", tickfont=dict(color="#212529"), showgrid=True,
                           gridcolor="#F0F0F0", showticklabels=False),
                coloraxis_showscale=False,
                plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
                margin=dict(t=10, b=10, l=10, r=10), height=320,
                font=dict(color="#212529"),
            )
            st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False}, theme=None)
        else:
            st.caption("Sem dados de entrada.")

    with kc2:
        st.markdown('<p class="sec-title">Distribuição de Conclusões por Dia da Semana</p>',
                    unsafe_allow_html=True)
        if not con_validas.empty:
            con_dow = con_validas["conclusao"].dt.dayofweek.value_counts().sort_index().reset_index()
            con_dow.columns = ["dow", "qtd"]
            con_dow["Dia"] = con_dow["dow"].map(lambda x: _dias_pt[x])
            fig2 = px.bar(
                con_dow, x="Dia", y="qtd", text="qtd",
                color="qtd",
                color_continuous_scale=["#C8F5C8", "#27AE60", "#145A32"],
            )
            fig2.update_traces(
                textposition="outside", cliponaxis=False, marker_line_width=0,
                hovertemplate="<b>%{x}</b><br>Conclusões: %{y}<extra></extra>",
            )
            fig2.update_layout(
                xaxis=dict(title="", tickfont=dict(color="#212529"), showgrid=False,
                           categoryorder="array", categoryarray=_dias_pt),
                yaxis=dict(title="Qtd", tickfont=dict(color="#212529"), showgrid=True,
                           gridcolor="#F0F0F0", showticklabels=False),
                coloraxis_showscale=False,
                plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
                margin=dict(t=10, b=10, l=10, r=10), height=320,
                font=dict(color="#212529"),
            )
            st.plotly_chart(fig2, use_container_width=True, config={"displayModeBar": False}, theme=None)
        else:
            st.caption("Sem dados de conclusão.")

    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    st.markdown('<p class="sec-title">Volume Mensal de Entradas e Conclusões</p>',
                unsafe_allow_html=True)

    ent_mes = (
        ent_validas.groupby(ent_validas["entrada"].dt.to_period("M"))
        .size().rename("Entradas")
    ) if not ent_validas.empty else pd.Series(dtype=int)
    con_mes = (
        con_validas.groupby(con_validas["conclusao"].dt.to_period("M"))
        .size().rename("Conclusões")
    ) if not con_validas.empty else pd.Series(dtype=int)

    dm_mes = pd.concat([ent_mes, con_mes], axis=1).fillna(0).reset_index()
    dm_mes.columns = ["Período", "Entradas", "Conclusões"]
    dm_mes["Período"] = dm_mes["Período"].astype(str)
    dm_mes = dm_mes.sort_values("Período")

    if not dm_mes.empty:
        fig3 = px.bar(
            dm_mes, x="Período", y=["Entradas", "Conclusões"], barmode="group",
            color_discrete_map={"Entradas": "#CC0000", "Conclusões": "#27AE60"},
        )
        fig3.update_layout(
            xaxis=dict(title="", tickfont=dict(color="#212529"), showgrid=False),
            yaxis=dict(title="Qtd", tickfont=dict(color="#212529"), showgrid=True,
                       gridcolor="#F0F0F0"),
            legend=dict(title="", orientation="h", y=1.05, x=0, font=dict(color="#212529")),
            plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
            margin=dict(t=20, b=10, l=10, r=10), height=300,
            font=dict(color="#212529"),
        )
        st.plotly_chart(fig3, use_container_width=True, config={"displayModeBar": False}, theme=None)


def _subtab_f2_gargalos(df_dem: pd.DataFrame) -> None:
    if df_dem.empty:
        st.info("Nenhum dado de demandas disponível.", icon="ℹ️")
        return

    pendentes = df_dem[df_dem["conclusao"].isna()].copy()

    if pendentes.empty:
        st.success("Nenhuma demanda pendente (sem conclusão) na seleção atual.", icon="✅")
        return

    kc1, kc2, kc3 = st.columns(3, gap="small")
    with kc1:
        st.metric("Demandas Pendentes", f"{len(pendentes):,}".replace(",", "."))
    with kc2:
        st.metric("Processos Afetados", f"{pendentes['processo_base'].nunique():,}".replace(",", "."))
    with kc3:
        st.metric("Procuradores com Pendências", f"{pendentes['procurador'].nunique():,}".replace(",", "."))

    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    st.markdown('<p class="sec-title">Tipos de Demanda Pendentes (Top 15)</p>',
                unsafe_allow_html=True)

    agg_tipo = (
        pendentes.groupby("demanda").size()
        .rename("Pendentes").reset_index()
        .nlargest(15, "Pendentes").sort_values("Pendentes")
    )
    fig = px.bar(
        agg_tipo, x="Pendentes", y="demanda", orientation="h", text="Pendentes",
        color="Pendentes",
        color_continuous_scale=["#FFE0B2", "#E07B00", "#7B3800"],
    )
    fig.update_traces(
        textposition="outside", cliponaxis=False, marker_line_width=0,
        hovertemplate="<b>%{y}</b><br>Pendentes: %{x}<extra></extra>",
    )
    fig.update_layout(
        xaxis=dict(showgrid=False, showticklabels=False, zeroline=False, title=""),
        yaxis=dict(showgrid=False, title="", tickfont=dict(size=10, color="#212529"),
                   automargin=True),
        coloraxis_showscale=False,
        plot_bgcolor="#FFFFFF", paper_bgcolor="#FFFFFF",
        margin=dict(t=10, b=10, l=5, r=55), height=420,
        font=dict(color="#212529"),
    )
    st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False}, theme=None)

    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    st.markdown('<p class="sec-title">Lista de Demandas Pendentes</p>', unsafe_allow_html=True)

    cols_pend = [c for c in ["processo_base", "procurador", "demanda", "entrada", "horas", "origem"]
                 if c in pendentes.columns]
    pend_tab = pendentes[cols_pend].copy()
    if "entrada" in pend_tab.columns:
        pend_tab["entrada"] = pend_tab["entrada"].dt.strftime("%d/%m/%Y").fillna("—")
    pend_tab = pend_tab.rename(columns={
        "processo_base": "Processo", "procurador": "Procurador",
        "demanda": "Demanda", "entrada": "Entrada", "horas": "Horas", "origem": "Origem",
    }).sort_values("Horas", ascending=False)

    st.dataframe(
        pend_tab, use_container_width=True, hide_index=True, height=400,
        column_config={
            "Processo":   st.column_config.TextColumn(width="medium"),
            "Procurador": st.column_config.TextColumn(width="medium"),
            "Demanda":    st.column_config.TextColumn(width="large"),
            "Entrada":    st.column_config.TextColumn(width="small"),
            "Horas":      st.column_config.NumberColumn(format="%.2f", width="small"),
            "Origem":     st.column_config.TextColumn(width="small"),
        },
    )
    st.markdown(
        f"<small style='color:#ADB5BD'>{len(pend_tab):,} demandas pendentes</small>".replace(",", "."),
        unsafe_allow_html=True,
    )


def _subtab_f2_procuradores(df_dem: pd.DataFrame) -> None:
    if df_dem.empty:
        st.info("Nenhum dado de demandas disponível.", icon="ℹ️")
        return

    concluidas = (
        df_dem[df_dem["conclusao"].notna()]
        .groupby("procurador").size().rename("Demandas Concluídas")
    )
    pendentes = (
        df_dem[df_dem["conclusao"].isna()]
        .groupby("procurador").size().rename("Demandas Pendentes")
    )

    tab = pd.concat([concluidas, pendentes], axis=1).fillna(0).reset_index()
    tab.columns = ["Procurador", "Demandas Concluídas", "Demandas Pendentes"]
    tab["Demandas Concluídas"] = tab["Demandas Concluídas"].astype(int)
    tab["Demandas Pendentes"]  = tab["Demandas Pendentes"].astype(int)
    tab["Total Geral"]         = tab["Demandas Concluídas"] + tab["Demandas Pendentes"]
    tab = tab.sort_values("Total Geral", ascending=False)

    c1, c2, c3 = st.columns(3, gap="small")
    with c1:
        st.metric("Procuradores", f"{len(tab):,}".replace(",", "."))
    with c2:
        st.metric("Total Concluídas", f"{int(tab['Demandas Concluídas'].sum()):,}".replace(",", "."))
    with c3:
        st.metric("Total Pendentes", f"{int(tab['Demandas Pendentes'].sum()):,}".replace(",", "."))

    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    st.markdown('<p class="sec-title">Demandas por Procurador</p>', unsafe_allow_html=True)

    st.dataframe(
        tab[["Procurador", "Demandas Concluídas", "Demandas Pendentes", "Total Geral"]],
        use_container_width=True, hide_index=True, height=560,
        column_config={
            "Procurador":           st.column_config.TextColumn(width="large"),
            "Demandas Concluídas":  st.column_config.NumberColumn(format="%d", width="small"),
            "Demandas Pendentes":   st.column_config.NumberColumn(format="%d", width="small"),
            "Total Geral":          st.column_config.NumberColumn(format="%d", width="small"),
        },
    )
    st.markdown(
        f"<small style='color:#ADB5BD'>{len(tab):,} procuradores</small>".replace(",", "."),
        unsafe_allow_html=True,
    )


def _subtab_f2_timeline(df: pd.DataFrame, df_dem_full: pd.DataFrame) -> None:
    if df.empty:
        st.warning("Nenhum processo corresponde à seleção. Ajuste os filtros.", icon="⚠️")
        return
    if df_dem_full.empty:
        st.info(
            "Nenhum dado de demandas disponível. "
            "Faça o upload do Relatório de Demandas na aba de atualização.",
            icon="ℹ️",
        )
        return

    st.markdown('<p class="sec-title">Histórico de Demandas por Processo</p>', unsafe_allow_html=True)
    st.markdown(
        '<p class="sec-sub">Selecione um processo para visualizar sua linha do tempo de demandas.</p>',
        unsafe_allow_html=True,
    )

    processos_opcoes = sorted(df["processo"].dropna().tolist())
    proc_sel = st.selectbox(
        "Processo",
        options=processos_opcoes,
        key="andamento_proc_sel",
        label_visibility="collapsed",
    )

    if proc_sel:
        hist = df_dem_full[df_dem_full["processo_base"] == proc_sel].copy()
        hist = hist.sort_values("conclusao", ascending=False, na_position="last")

        if hist.empty:
            st.info(f"Nenhuma demanda registrada para o processo **{proc_sel}**.", icon="ℹ️")
            return

        proc_info = df[df["processo"] == proc_sel]
        if not proc_info.empty:
            pi = proc_info.iloc[0]
            c1, c2, c3, c4 = st.columns(4, gap="small")
            with c1:
                st.metric("Status", str(pi.get("status_exito", "—")))
            with c2:
                st.metric("Procurador Atual", str(pi.get("procurador", "—")))
            with c3:
                st.metric("Total de Horas", f"{float(pi.get('total_horas', 0) or 0):.1f}")
            with c4:
                st.metric("Nº de Demandas", f"{len(hist):,}".replace(",", "."))

        st.markdown('<hr class="divider">', unsafe_allow_html=True)

        hist_disp = hist[[
            "entrada", "conclusao", "demanda", "procurador", "origem", "horas", "publicou"
        ]].copy()
        hist_disp.columns = [
            "Data Entrada", "Data Conclusão", "Demanda", "Procurador", "Origem", "Horas", "Publicou"
        ]
        hist_disp["Data Entrada"]   = hist_disp["Data Entrada"].dt.strftime("%d/%m/%Y").fillna("—")
        hist_disp["Data Conclusão"] = hist_disp["Data Conclusão"].dt.strftime("%d/%m/%Y").fillna("—")

        st.dataframe(
            hist_disp, use_container_width=True, hide_index=True, height=420,
            column_config={
                "Data Entrada":   st.column_config.TextColumn(width="small"),
                "Data Conclusão": st.column_config.TextColumn(width="small"),
                "Demanda":        st.column_config.TextColumn(width="large"),
                "Procurador":     st.column_config.TextColumn(width="medium"),
                "Origem":         st.column_config.TextColumn(width="small"),
                "Horas":          st.column_config.NumberColumn(format="%.2f", width="small"),
                "Publicou":       st.column_config.TextColumn(width="small"),
            },
        )
        st.markdown(
            f"<small style='color:#ADB5BD'>{len(hist_disp)} demandas</small>",
            unsafe_allow_html=True,
        )


def _tab_upload() -> None:
    st.markdown('<p class="sec-title">Atualização da Base de Dados</p>', unsafe_allow_html=True)
    st.markdown(
        '<p class="sec-sub">Importe os dois relatórios exportados do sistema de gestão processual. '
        'O núcleo é identificado automaticamente a partir dos arquivos.</p>',
        unsafe_allow_html=True,
    )

    col_form, col_info = st.columns([3, 2], gap="large")

    grupo_opts = {
        "Ambiental e Imobiliário":      _NUCLEOS[:2],
        "Pessoal e Previdenciário":     _NUCLEOS[2:11],
        "Serviços Públicos e Residual": _NUCLEOS[11:],
    }

    with col_form:
        arquivo_processos = st.file_uploader(
            "📋 Relatório de Processos (Relatório Detalhado para Excel)",
            type=["csv", "txt"],
            key="upload_processos",
            help="Arquivo: 'Processos - Relatório Detalhado para Excel (Lista) (Mês Ano).txt'",
        )
        arquivo_demandas = st.file_uploader(
            "📋 Relatório de Demandas (Por unidade e procurador)",
            type=["csv", "txt"],
            key="upload_demandas",
            help="Arquivo: 'Demandas por unidade e procurador para Excel (Lista) (Mês Ano).txt'",
        )

        ambos_carregados = arquivo_processos is not None and arquivo_demandas is not None

        if not ambos_carregados and (arquivo_processos or arquivo_demandas):
            faltando = "Demandas" if arquivo_processos else "Processos"
            st.caption(f"⚠️ Aguardando arquivo de **{faltando}** para habilitar o processamento.")

        nucleo_detectado = None
        if ambos_carregados:
            nucleo_detectado = (
                _detectar_nucleo_arquivo(arquivo_processos)
                or _detectar_nucleo_arquivo(arquivo_demandas)
            )

        st.markdown("<div style='margin-top:.8rem'>", unsafe_allow_html=True)
        if nucleo_detectado:
            st.success(
                f"🔎 **Núcleo identificado automaticamente:** {nucleo_detectado}",
                icon="✅",
            )
            nucleo_final = nucleo_detectado
            usar_manual = st.checkbox(
                "Núcleo incorreto? Selecionar manualmente",
                key="upload_usar_manual",
            )
            if usar_manual:
                grupo_padrao = _grupo_do_nucleo(nucleo_detectado, grupo_opts) or next(iter(grupo_opts))
                grupos_lista = list(grupo_opts.keys())
                grupo_sel = st.selectbox(
                    "Grupo",
                    options=grupos_lista,
                    index=grupos_lista.index(grupo_padrao),
                    key="upload_grupo",
                )
                lista_nucleos = grupo_opts[grupo_sel]
                idx_padrao = lista_nucleos.index(nucleo_detectado) if nucleo_detectado in lista_nucleos else 0
                nucleo_final = st.selectbox(
                    "Núcleo", options=lista_nucleos, index=idx_padrao, key="upload_nucleo",
                )
        elif ambos_carregados:
            st.warning(
                "Não foi possível identificar o núcleo automaticamente a partir dos arquivos. "
                "Selecione manualmente abaixo.",
                icon="⚠️",
            )
            grupo_sel = st.selectbox("Grupo", options=list(grupo_opts.keys()), key="upload_grupo")
            nucleo_final = st.selectbox("Núcleo", options=grupo_opts[grupo_sel], key="upload_nucleo")
        else:
            st.caption(
                "Envie os dois arquivos para identificar o núcleo automaticamente, "
                "ou selecione manualmente abaixo."
            )
            grupo_sel = st.selectbox("Grupo", options=list(grupo_opts.keys()), key="upload_grupo")
            nucleo_final = st.selectbox("Núcleo", options=grupo_opts[grupo_sel], key="upload_nucleo")
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<div style='margin-top:.8rem'>", unsafe_allow_html=True)
        processar = st.button(
            "⬆️  Processar e Salvar Base",
            type="primary",
            use_container_width=True,
            disabled=not ambos_carregados,
        )
        st.markdown("</div>", unsafe_allow_html=True)

        if processar:
            with st.spinner(f"Processando base do {nucleo_final}..."):
                try:
                    qtd_proc, qtd_dem = processar_upload(
                        arquivo_processos, arquivo_demandas, nucleo_final
                    )
                    st.success(
                        f"**{qtd_proc:,} processos** e **{qtd_dem:,} demandas** salvos "
                        f"com sucesso para **{nucleo_final}**!".replace(",", "."),
                        icon="✅",
                    )
                    st.cache_data.clear()
                    st.rerun()
                except KeyError as e:
                    st.error(f"Coluna não encontrada no arquivo: {e}", icon="❌")
                except ValueError as e:
                    st.error(str(e), icon="❌")
                except Exception as e:
                    st.error(f"Erro inesperado ao processar o arquivo: {e}", icon="❌")

    with col_info:
        st.info(
            "**Formato esperado**\n\n"
            "**Arquivo 1 — Processos:**\n"
            "- 'Relatório Detalhado para Excel (Lista)'\n"
            "- CSV separado por vírgula (`.txt` ou `.csv`)\n\n"
            "**Arquivo 2 — Demandas:**\n"
            "- 'Demandas por unidade e procurador para Excel (Lista)'\n"
            "- CSV separado por vírgula (`.txt` ou `.csv`)\n\n"
            "**O que acontece ao processar:**\n\n"
            "- Processos novos são **inseridos**\n"
            "- Processos existentes têm valor, tramitação e situação **atualizados**\n"
            "- Demandas antigas do núcleo são **substituídas** pelas novas\n"
            "- Status de êxito, procurador e horas são **derivados das demandas**",
            icon="ℹ️",
        )

        st.markdown(
            "<p style='font-size:.62rem;font-weight:800;text-transform:uppercase;"
            "letter-spacing:2px;color:#CC0000;margin-top:1.2rem;margin-bottom:.4rem'>"
            "Histórico de Uploads</p>",
            unsafe_allow_html=True,
        )
        if _db.db_exists():
            try:
                con = _db.connect()
                hist = _db.read_sql(
                    "SELECT data_upload, nome_arquivo, quantidade_registros_processados, nucleo "
                    "FROM controle_uploads ORDER BY id DESC LIMIT 10",
                    con,
                )
                con.close()
                if not hist.empty:
                    hist.columns = ["Data/Hora", "Arquivo", "Registros", "Núcleo"]
                    hist["Data/Hora"] = hist["Data/Hora"].str[:16].str.replace("T", " ")
                    hist["Núcleo"] = hist["Núcleo"].fillna("—")
                    st.dataframe(hist, use_container_width=True, hide_index=True, height=240)
                else:
                    st.caption("Nenhum upload registrado ainda.")
            except Exception:
                st.caption("Histórico indisponível.")
        else:
            st.caption("Banco de dados ainda não criado.")

    # ── Exclusão de dados ────────────────────────────────────────────────
    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    st.markdown('<p class="sec-title">🗑️ Excluir Dados de um Núcleo</p>', unsafe_allow_html=True)
    st.markdown(
        '<p class="sec-sub">Remove permanentemente processos, demandas e histórico de uploads '
        'de um núcleo específico. Esta ação não pode ser desfeita.</p>',
        unsafe_allow_html=True,
    )

    nucleos_existentes = _nucleos_carregados()

    if nucleos_existentes:
        col_excl, col_resumo = st.columns([2, 1], gap="medium")

        with col_excl:
            nucleo_excluir = st.selectbox(
                "Núcleo a excluir",
                options=nucleos_existentes,
                key="excluir_nucleo_sel",
            )
            resumo = _resumo_nucleo(nucleo_excluir)

            confirmar = st.checkbox(
                f"Confirmo que desejo excluir permanentemente todos os dados de "
                f"**{nucleo_excluir}** ({resumo['processos']:,} processos, "
                f"{resumo['demandas']:,} demandas)".replace(",", "."),
                key="excluir_nucleo_confirmar",
            )
            excluir = st.button(
                "🗑️  Excluir Dados Permanentemente",
                type="primary",
                use_container_width=True,
                disabled=not confirmar,
            )

            if excluir:
                try:
                    resultado = excluir_dados_nucleo(nucleo_excluir)
                    st.success(
                        f"Dados de **{nucleo_excluir}** excluídos: "
                        f"{resultado['processos']:,} processos, "
                        f"{resultado['demandas']:,} demandas, "
                        f"{resultado['uploads']:,} registros de upload.".replace(",", "."),
                        icon="✅",
                    )
                    st.session_state.pop("excluir_nucleo_confirmar", None)
                    st.cache_data.clear()
                    st.rerun()
                except Exception as e:
                    st.error(f"Erro ao excluir dados: {e}", icon="❌")

        with col_resumo:
            st.metric("Processos", f"{resumo['processos']:,}".replace(",", "."))
            st.metric("Demandas", f"{resumo['demandas']:,}".replace(",", "."))
            st.metric("Registros de Upload", f"{resumo['uploads']:,}".replace(",", "."))
    else:
        st.caption(
            "Nenhum processo com núcleo identificado no momento "
            "(dados podem ter sido carregados sem marcação de núcleo)."
        )

    # ── Reset total ──────────────────────────────────────────────────────
    st.markdown('<hr class="divider">', unsafe_allow_html=True)
    st.markdown(
        '<p class="sec-title" style="color:#8B0000">⚠️ Reset Total — Excluir Todos os Dados</p>',
        unsafe_allow_html=True,
    )
    st.markdown(
        '<p class="sec-sub">Remove permanentemente <b>todos os processos, demandas e histórico de '
        'uploads, de todos os núcleos (inclusive sem núcleo identificado)</b>, devolvendo o sistema '
        'ao estado inicial. Esta ação não pode ser desfeita.</p>',
        unsafe_allow_html=True,
    )

    resumo_total = _resumo_geral()

    if resumo_total["processos"] == 0 and resumo_total["demandas"] == 0 and resumo_total["uploads"] == 0:
        st.caption("Não há dados no banco para excluir.")
        return

    col_reset, col_resumo_total = st.columns([2, 1], gap="medium")

    with col_reset:
        st.caption(
            f"Serão excluídos **{resumo_total['processos']:,} processos**, "
            f"**{resumo_total['demandas']:,} demandas** e "
            f"**{resumo_total['uploads']:,} registros de upload**.".replace(",", ".")
        )
        frase_digitada = st.text_input(
            'Para confirmar, digite **RESETAR TUDO** no campo abaixo',
            key="reset_total_frase",
            placeholder="RESETAR TUDO",
        )
        reset_habilitado = frase_digitada.strip().upper() == "RESETAR TUDO"
        resetar = st.button(
            "🗑️  Excluir TODOS os Dados (Reset Total)",
            type="primary",
            use_container_width=True,
            disabled=not reset_habilitado,
        )

        if resetar:
            try:
                resultado_total = excluir_todos_dados()
                st.success(
                    f"Reset concluído: {resultado_total['processos']:,} processos, "
                    f"{resultado_total['demandas']:,} demandas e "
                    f"{resultado_total['uploads']:,} registros de upload excluídos.".replace(",", "."),
                    icon="✅",
                )
                st.session_state.pop("reset_total_frase", None)
                st.cache_data.clear()
                st.rerun()
            except Exception as e:
                st.error(f"Erro ao executar o reset total: {e}", icon="❌")

    with col_resumo_total:
        st.metric("Núcleos Identificados", f"{len(nucleos_existentes):,}".replace(",", "."))
        st.metric("Processos (total)", f"{resumo_total['processos']:,}".replace(",", "."))
        st.metric("Demandas (total)", f"{resumo_total['demandas']:,}".replace(",", "."))

# =============================================================================
# GERAÇÃO DE RELATÓRIOS WORD
# =============================================================================

def _substituir_placeholders(doc, mapa: dict) -> None:
    """Substitui {{PLACEHOLDER}} em todos os parágrafos e tabelas do documento."""
    def _proc_para(para):
        for run in para.runs:
            for chave, valor in mapa.items():
                if chave in run.text:
                    run.text = run.text.replace(chave, str(valor))

    for para in doc.paragraphs:
        _proc_para(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _proc_para(para)
                for nested in cell.tables:
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            for npara in ncell.paragraphs:
                                _proc_para(npara)


def _gerar_relatorio(template_path: Path, mapa: dict) -> bytes | None:
    """Abre o template .docx, substitui os placeholders e retorna bytes."""
    try:
        from docx import Document  # type: ignore[import]
    except ImportError:
        st.error("Pacote 'python-docx' não instalado. Execute: pip install python-docx", icon="❌")
        return None
    if not template_path.exists():
        st.warning(f"Template não encontrado: {template_path.name}", icon="⚠️")
        return None
    try:
        import io as _io
        doc = Document(str(template_path))
        _substituir_placeholders(doc, mapa)
        buf = _io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
    except Exception as exc:
        st.error(f"Erro ao gerar relatório: {exc}", icon="❌")
        return None


def _build_mapa_frente1(df: pd.DataFrame, filtros_sidebar: dict) -> dict:
    """Monta o dicionário de substituição para o relatório da Frente 1."""
    hoje = date.today().strftime("%d/%m/%Y")

    # ── Metadados ─────────────────────────────────────────────────────────────
    nucleo_label = ", ".join(filtros_sidebar.get("nucleo", [])) or "Todos os núcleos"
    periodo_label = "—"
    if not df.empty and "ajuizamento" in df.columns:
        datas = df["ajuizamento"].dropna()
        if not datas.empty:
            periodo_label = (
                f"{datas.min().strftime('%m/%Y')} a {datas.max().strftime('%m/%Y')}"
            )
    filtros_txt_parts = []
    if filtros_sidebar.get("nucleo"):
        filtros_txt_parts.append(f"Núcleo: {nucleo_label}")
    if filtros_sidebar.get("procurador"):
        filtros_txt_parts.append(
            "Procurador: " + ", ".join(filtros_sidebar["procurador"])
        )
    filtros_label = "; ".join(filtros_txt_parts) or "Sem filtros adicionais"

    # ── KPIs ──────────────────────────────────────────────────────────────────
    kpi = _kpis(df) if not df.empty else {
        "total": 0, "em_risco": 0.0, "economia": 0.0,
        "taxa": 0.0, "vitorias": 0, "perdas": 0, "decididos": 0,
    }
    valor_medio = (df["valor"].mean() if not df.empty else 0.0)
    n_comarcas  = (df["comarca_limpa"].nunique() if not df.empty and "comarca_limpa" in df.columns else 0)
    n_andamento = int((df["status_exito"] == "Em Andamento").sum()) if not df.empty else 0

    # ── Teses (top 5 por volume) ───────────────────────────────────────────────
    if not df.empty and "assunto_label" in df.columns:
        teses_df = (
            df.groupby("assunto_label")
            .agg(
                total    = ("processo",      "count"),
                vitorias = ("status_exito",  lambda x: (x == "Vitória").sum()),
                perdas   = ("status_exito",  lambda x: (x == "Perda").sum()),
                valor    = ("valor",         "sum"),
            )
            .assign(
                em_andamento = lambda d: d["total"] - d["vitorias"] - d["perdas"],
                taxa         = lambda d: (
                    (d["vitorias"] / (d["vitorias"] + d["perdas"]) * 100)
                    .where((d["vitorias"] + d["perdas"]) > 0)
                    .fillna(0).round(1)
                ),
            )
            .sort_values("total", ascending=False)
            .head(5)
            .reset_index()
        )
    else:
        teses_df = pd.DataFrame()

    def _t(i, col, fmt=None):
        if i < len(teses_df):
            val = teses_df.iloc[i][col]
            return _brl(val) if fmt == "brl" else (f"{val:.1f}%" if fmt == "pct" else str(val))
        return "—"

    # ── Êxitos por tese (reutiliza teses_df) ──────────────────────────────────
    def _e(i, col, fmt=None):
        return _t(i, col, fmt)

    # ── Comarcas (top 10) ─────────────────────────────────────────────────────
    if not df.empty and "comarca_limpa" in df.columns:
        com_df = (
            df.groupby("comarca_limpa")
            .agg(
                total    = ("processo",     "count"),
                vitorias = ("status_exito", lambda x: (x == "Vitória").sum()),
                perdas   = ("status_exito", lambda x: (x == "Perda").sum()),
                valor    = ("valor",        "sum"),
            )
            .assign(
                taxa = lambda d: (
                    (d["vitorias"] / (d["vitorias"] + d["perdas"]) * 100)
                    .where((d["vitorias"] + d["perdas"]) > 0)
                    .fillna(0).round(1)
                )
            )
            .sort_values("total", ascending=False)
            .head(10)
            .reset_index()
        )
    else:
        com_df = pd.DataFrame()

    def _c(i, col, fmt=None):
        if i < len(com_df):
            val = com_df.iloc[i][col]
            return _brl(val) if fmt == "brl" else (f"{val:.1f}%" if fmt == "pct" else str(val))
        return "—"

    # ── Linha do tempo (últimos 6 meses por ajuizamento) ──────────────────────
    if not df.empty and "ajuizamento" in df.columns:
        tmp = df.copy()
        tmp["mes"] = tmp["ajuizamento"].dt.to_period("M")
        tempo_df = (
            tmp.groupby("mes")
            .agg(
                novos    = ("processo",     "count"),
                decididos= ("status_exito", lambda x: ((x=="Vitória")|(x=="Perda")).sum()),
                vitorias = ("status_exito", lambda x: (x=="Vitória").sum()),
                perdas   = ("status_exito", lambda x: (x=="Perda").sum()),
                andamento= ("status_exito", lambda x: (x=="Em Andamento").sum()),
                valor    = ("valor",        "sum"),
            )
            .tail(6)
            .reset_index()
        )
        tempo_df["mes_str"] = tempo_df["mes"].astype(str)
    else:
        tempo_df = pd.DataFrame()

    def _tmp(i, col, fmt=None):
        if i < len(tempo_df):
            val = tempo_df.iloc[i][col]
            return _brl(val) if fmt == "brl" else str(val)
        return "—"

    # ── Detalhamento (top 10 por valor) ───────────────────────────────────────
    if not df.empty:
        det_df = df.nlargest(10, "valor").reset_index(drop=True)
    else:
        det_df = pd.DataFrame()

    def _d(i, col, fmt=None):
        if i < len(det_df):
            val = det_df.iloc[i].get(col, "—")
            if pd.isna(val):
                return "—"
            return _brl(float(val)) if fmt == "brl" else str(val)
        return "—"

    def _d_date(i, col):
        if i < len(det_df):
            val = det_df.iloc[i].get(col)
            if pd.isna(val) if val is not None else True:
                return "—"
            try:
                return pd.Timestamp(val).strftime("%d/%m/%Y")
            except Exception:
                return str(val)
        return "—"

    # ── Mapa final ────────────────────────────────────────────────────────────
    mapa: dict = {
        "{{DATA_GERACAO}}":     hoje,
        "{{NUCLEO}}":           nucleo_label,
        "{{PERIODO}}":          periodo_label,
        "{{FILTROS_APLICADOS}}": filtros_label,

        # KPIs
        "{{TOTAL_PROCESSOS}}":  f"{kpi['total']:,}".replace(",", "."),
        "{{QTD_EM_ANDAMENTO}}": f"{n_andamento:,}".replace(",", "."),
        "{{VALOR_EM_RISCO}}":   _brl(kpi["em_risco"]),
        "{{VALOR_ECONOMIA}}":   _brl(kpi["economia"]),
        "{{QTD_VITORIAS}}":     str(kpi["vitorias"]),
        "{{QTD_PERDAS}}":       str(kpi["perdas"]),
        "{{QTD_DECIDIDOS}}":    str(kpi["decididos"]),
        "{{TAXA_EXITO}}":       f"{kpi['taxa']:.1f}%",
        "{{VALOR_MEDIO}}":      _brl(valor_medio),
        "{{QTD_COMARCAS}}":     str(n_comarcas),
    }

    # Teses 1–5
    for i in range(5):
        n = i + 1
        mapa[f"{{{{TESE_{n}}}}}"]  = _t(i, "assunto_label")
        mapa[f"{{{{N{n}}}}}"]       = _t(i, "total")
        mapa[f"{{{{V{n}}}}}"]       = _t(i, "vitorias")
        mapa[f"{{{{P{n}}}}}"]       = _t(i, "perdas")
        mapa[f"{{{{A{n}}}}}"]       = _t(i, "em_andamento")
        mapa[f"{{{{T{n}}}}}"]       = _t(i, "taxa", "pct")
        mapa[f"{{{{VAL{n}}}}}"]     = _t(i, "valor", "brl")
        # Êxitos (mesma fonte)
        mapa[f"{{{{E{n}_N}}}}"]     = _t(i, "assunto_label")
        mapa[f"{{{{E{n}_P}}}}"]     = _t(i, "taxa", "pct")
        mapa[f"{{{{E{n}_VAL}}}}"]   = _t(i, "valor", "brl")
        mapa[f"{{{{E{n}_T}}}}"]     = _t(i, "total")
        mapa[f"{{{{E{n}_PD}}}}"]    = str(
            int(teses_df.iloc[i]["vitorias"] + teses_df.iloc[i]["perdas"])
            if i < len(teses_df) else "—"
        )

    # Totais teses
    if not teses_df.empty:
        mapa["{{NTOT}}"]   = str(int(teses_df["total"].sum()))
        mapa["{{VTOT}}"]   = str(int(teses_df["vitorias"].sum()))
        mapa["{{PTOT}}"]   = str(int(teses_df["perdas"].sum()))
        mapa["{{ATOT}}"]   = str(int(teses_df["em_andamento"].sum()))
        mapa["{{TTOT}}"]   = "—"
        mapa["{{VALTOT}}"] = _brl(teses_df["valor"].sum())
        mapa["{{ETOT}}"]   = str(int(teses_df["total"].sum()))
        mapa["{{ETOT_VAL}}"] = _brl(teses_df["valor"].sum())
        mapa["{{ETOT_T}}"] = str(int(teses_df["total"].sum()))
        mapa["{{ETOT_PD}}"]= str(int((teses_df["vitorias"] + teses_df["perdas"]).sum()))
    else:
        for k in ("{{NTOT}}", "{{VTOT}}", "{{PTOT}}", "{{ATOT}}", "{{TTOT}}", "{{VALTOT}}",
                  "{{ETOT}}", "{{ETOT_VAL}}", "{{ETOT_T}}", "{{ETOT_PD}}"):
            mapa[k] = "—"

    # Comarcas 1–10
    for i in range(10):
        n = i + 1
        mapa[f"{{{{COMARCA_{n}}}}}"] = _c(i, "comarca_limpa")
        mapa[f"{{{{C{n}_T}}}}"]      = _c(i, "total")
        mapa[f"{{{{C{n}_V}}}}"]      = _c(i, "vitorias")
        mapa[f"{{{{C{n}_P}}}}"]      = _c(i, "perdas")
        mapa[f"{{{{C{n}_TX}}}}"]     = _c(i, "taxa", "pct")
        mapa[f"{{{{C{n}_VAL}}}}"]    = _c(i, "valor", "brl")

    # Linha do tempo 1–6
    for i in range(6):
        n = i + 1
        mapa[f"{{{{MES_{n}}}}}"]    = _tmp(i, "mes_str")
        mapa[f"{{{{TMP{n}_N}}}}"]   = _tmp(i, "novos")
        mapa[f"{{{{TMP{n}_D}}}}"]   = _tmp(i, "decididos")
        mapa[f"{{{{TMP{n}_V}}}}"]   = _tmp(i, "vitorias")
        mapa[f"{{{{TMP{n}_P}}}}"]   = _tmp(i, "perdas")
        mapa[f"{{{{TMP{n}_A}}}}"]   = _tmp(i, "andamento")
        mapa[f"{{{{TMP{n}_VAL}}}}"] = _tmp(i, "valor", "brl")

    # Detalhamento 1–10
    for i in range(10):
        n = i + 1
        mapa[f"{{{{PROC_{n}}}}}"]   = _d(i, "processo")
        mapa[f"{{{{DAJ_{n}}}}}"]    = _d_date(i, "ajuizamento")
        mapa[f"{{{{TESE_D_{n}}}}}"] = _d(i, "assunto_label")
        mapa[f"{{{{COM_{n}}}}}"]    = _d(i, "comarca_limpa")
        mapa[f"{{{{EST_{n}}}}}"]    = _d(i, "status_exito")
        mapa[f"{{{{RES_{n}}}}}"]    = _d(i, "status_exito")
        mapa[f"{{{{VAL_D_{n}}}}}"]  = _d(i, "valor", "brl")

    return mapa


def _build_mapa_frente2(df_dem: pd.DataFrame, filtros_sidebar: dict) -> dict:
    """Monta o dicionário de substituição para o relatório da Frente 2."""
    hoje = date.today().strftime("%d/%m/%Y")

    # ── Metadados ─────────────────────────────────────────────────────────────
    nucleo_label = ", ".join(filtros_sidebar.get("nucleo", [])) or "Todos os núcleos"
    periodo_label = "—"
    if not df_dem.empty:
        datas = df_dem["conclusao"].dropna()
        if not datas.empty:
            periodo_label = (
                f"{datas.min().strftime('%m/%Y')} a {datas.max().strftime('%m/%Y')}"
            )
    filtros_txt_parts = []
    if filtros_sidebar.get("nucleo"):
        filtros_txt_parts.append(f"Núcleo: {nucleo_label}")
    if filtros_sidebar.get("procurador"):
        filtros_txt_parts.append(
            "Procurador: " + ", ".join(filtros_sidebar["procurador"])
        )
    filtros_label = "; ".join(filtros_txt_parts) or "Sem filtros adicionais"

    # ── KPIs ──────────────────────────────────────────────────────────────────
    total_dem  = len(df_dem)
    conc_mask  = df_dem["conclusao"].notna() if not df_dem.empty else pd.Series(dtype=bool)
    pend_mask  = df_dem["conclusao"].isna()  if not df_dem.empty else pd.Series(dtype=bool)
    dem_conc   = int(conc_mask.sum())
    dem_pend   = int(pend_mask.sum())
    perc_conc  = f"{dem_conc/total_dem*100:.1f}%" if total_dem > 0 else "—"
    total_h    = df_dem["horas"].sum() if not df_dem.empty else 0.0
    media_h    = df_dem["horas"].mean() if not df_dem.empty else 0.0
    qtd_proc   = df_dem["procurador"].nunique() if not df_dem.empty else 0
    qtd_pub    = int((df_dem["publicou"] == "Sim").sum()) if not df_dem.empty and "publicou" in df_dem.columns else 0
    qtd_tipos  = df_dem["demanda"].nunique() if not df_dem.empty and "demanda" in df_dem.columns else 0

    # Média de dias (entrada → conclusao)
    if not df_dem.empty and "entrada" in df_dem.columns and "conclusao" in df_dem.columns:
        delta = (df_dem["conclusao"] - df_dem["entrada"]).dt.days.dropna()
        media_dias = f"{delta.mean():.0f}" if not delta.empty else "—"
    else:
        media_dias = "—"

    # ── Produtividade por procurador (top 10) ─────────────────────────────────
    if not df_dem.empty and "procurador" in df_dem.columns:
        prod_df = (
            df_dem.groupby("procurador")
            .agg(
                concluidas  = ("conclusao", lambda x: x.notna().sum()),
                pendentes   = ("conclusao", lambda x: x.isna().sum()),
                horas       = ("horas",     "sum"),
                publicacoes = ("publicou",  lambda x: (x == "Sim").sum()),
            )
            .assign(
                total       = lambda d: d["concluidas"] + d["pendentes"],
                media_horas = lambda d: (d["horas"] / d["total"]).where(d["total"] > 0).fillna(0).round(1),
            )
            .sort_values("total", ascending=False)
            .head(10)
            .reset_index()
        )
    else:
        prod_df = pd.DataFrame()

    def _p(i, col, fmt=None):
        if i < len(prod_df):
            val = prod_df.iloc[i][col]
            return f"{val:.1f}" if fmt == "f1" else str(int(val)) if fmt == "int" else str(val)
        return "—"

    # ── Fluxo mensal (últimos 6 meses por conclusao) ───────────────────────────
    if not df_dem.empty and "conclusao" in df_dem.columns and "entrada" in df_dem.columns:
        conc_v = df_dem.dropna(subset=["conclusao"]).copy()
        ent_v  = df_dem.dropna(subset=["entrada"]).copy()
        conc_v["mes"] = conc_v["conclusao"].dt.to_period("M")
        ent_v["mes"]  = ent_v["entrada"].dt.to_period("M")
        conc_mes = conc_v.groupby("mes").size().rename("conc")
        ent_mes  = ent_v.groupby("mes").size().rename("ent")
        flux_df  = pd.concat([ent_mes, conc_mes], axis=1).fillna(0)
        flux_df["saldo"] = flux_df["conc"] - flux_df["ent"]
        flux_df["sacum"] = flux_df["saldo"].cumsum()
        flux_df["pend"]  = df_dem.groupby(
            df_dem["entrada"].dt.to_period("M")
        ).apply(lambda g: g["conclusao"].isna().sum(), include_groups=False).reindex(flux_df.index).fillna(0)
        flux_df = flux_df.tail(6).reset_index()
        flux_df["mes_str"] = flux_df["mes"].astype(str)
    else:
        flux_df = pd.DataFrame()

    def _fl(i, col):
        if i < len(flux_df):
            val = flux_df.iloc[i][col]
            return str(int(val)) if isinstance(val, float) else str(val)
        return "—"

    # ── Tipos de demanda (top 8) ───────────────────────────────────────────────
    if not df_dem.empty and "demanda" in df_dem.columns:
        tipo_df = (
            df_dem.groupby("demanda")
            .agg(
                n    = ("demanda",   "count"),
                pend = ("conclusao", lambda x: x.isna().sum()),
            )
            .sort_values("n", ascending=False)
            .head(8)
            .reset_index()
        )
    else:
        tipo_df = pd.DataFrame()

    def _td(i, col):
        if i < len(tipo_df):
            return str(tipo_df.iloc[i][col])
        return "—"

    # ── Sazonalidade (por mês do ano, usando entrada) ─────────────────────────
    _meses_abrev = ["JAN","FEV","MAR","ABR","MAI","JUN",
                    "JUL","AGO","SET","OUT","NOV","DEZ"]
    if not df_dem.empty and "entrada" in df_dem.columns:
        saz = df_dem.dropna(subset=["entrada"]).copy()
        saz["mes_num"] = saz["entrada"].dt.month
        saz_agg = (
            saz.groupby("mes_num")
            .agg(ent=("entrada","count"), horas=("horas","sum"))
            .reindex(range(1,13), fill_value=0)
        )
    else:
        saz_agg = pd.DataFrame({"ent": [0]*12, "horas": [0.0]*12}, index=range(1,13))

    # ── Gargalos — 10 demandas pendentes mais antigas ────────────────────────
    if not df_dem.empty and "entrada" in df_dem.columns:
        garg_df = (
            df_dem[df_dem["conclusao"].isna()]
            .dropna(subset=["entrada"])
            .copy()
        )
        garg_df["dias"] = (pd.Timestamp(date.today()) - garg_df["entrada"]).dt.days
        garg_df = garg_df.nlargest(10, "dias").reset_index(drop=True)
    else:
        garg_df = pd.DataFrame()

    def _g(i, col, fmt=None):
        if i < len(garg_df):
            val = garg_df.iloc[i].get(col, "—")
            if pd.isna(val) if not isinstance(val, str) else (val == ""):
                return "—"
            return val.strftime("%d/%m/%Y") if fmt == "date" and hasattr(val, "strftime") else str(val)
        return "—"

    # ── Histórico 12 meses ────────────────────────────────────────────────────
    if not df_dem.empty and "conclusao" in df_dem.columns:
        hist_c = df_dem.dropna(subset=["conclusao"]).copy()
        hist_e = df_dem.dropna(subset=["entrada"]).copy()
        hist_c["mes"] = hist_c["conclusao"].dt.to_period("M")
        hist_e["mes"] = hist_e["entrada"].dt.to_period("M")
        hist_conc = hist_c.groupby("mes").agg(
            conc=("conclusao","count"), horas=("horas","sum"),
            pub=("publicou", lambda x: (x=="Sim").sum()),
        )
        hist_ent = hist_e.groupby("mes").size().rename("ent")
        hist_df  = pd.concat([hist_ent, hist_conc], axis=1).fillna(0)
        hist_df["mh"]  = (hist_df["horas"] / hist_df["conc"].replace(0, pd.NA)).fillna(0).round(1)
        hist_df["tx"]  = (hist_df["conc"] / (hist_df["ent"] + hist_df["conc"]).replace(0, pd.NA) * 100).fillna(0).round(1)
        hist_df = hist_df.tail(12).reset_index()
        hist_df["mes_str"] = hist_df["mes"].astype(str)
    else:
        hist_df = pd.DataFrame()

    def _h(i, col, fmt=None):
        if i < len(hist_df):
            val = hist_df.iloc[i][col]
            return f"{val:.1f}" if fmt == "f1" else str(int(val)) if fmt == "int" else str(val)
        return "—"

    # ── Mapa final ────────────────────────────────────────────────────────────
    mapa: dict = {
        "{{DATA_GERACAO}}":      hoje,
        "{{NUCLEO}}":            nucleo_label,
        "{{PERIODO}}":           periodo_label,
        "{{FILTROS_APLICADOS}}": filtros_label,

        # KPIs
        "{{TOTAL_DEMANDAS}}":  f"{total_dem:,}".replace(",", "."),
        "{{DEMANDAS_CONC}}":   f"{dem_conc:,}".replace(",", "."),
        "{{PERC_CONC}}":       perc_conc,
        "{{DEMANDAS_PEND}}":   f"{dem_pend:,}".replace(",", "."),
        "{{TOTAL_HORAS}}":     f"{total_h:,.1f}".replace(",", "."),
        "{{MEDIA_HORAS}}":     f"{media_h:.1f}",
        "{{QTD_PROC}}":        str(qtd_proc),
        "{{QTD_PUBLICACOES}}": str(qtd_pub),
        "{{QTD_TIPOS}}":       str(qtd_tipos),
        "{{MEDIA_DIAS}}":      media_dias,
    }

    # Produtividade por procurador 1–10
    for i in range(10):
        n = i + 1
        mapa[f"{{{{PROC_{n}}}}}"]   = _p(i, "procurador")
        mapa[f"{{{{PC{n}_CONC}}}}"] = _p(i, "concluidas", "int")
        mapa[f"{{{{PC{n}_PEND}}}}"] = _p(i, "pendentes",  "int")
        mapa[f"{{{{PC{n}_TOT}}}}"]  = _p(i, "total",      "int")
        mapa[f"{{{{PC{n}_H}}}}"]    = _p(i, "horas",      "f1")
        mapa[f"{{{{PC{n}_MH}}}}"]   = _p(i, "media_horas","f1")
        mapa[f"{{{{PC{n}_PUB}}}}"]  = _p(i, "publicacoes","int")

    # Totais produtividade
    if not prod_df.empty:
        mapa["{{PTOT_CONC}}"] = str(int(prod_df["concluidas"].sum()))
        mapa["{{PTOT_PEND}}"] = str(int(prod_df["pendentes"].sum()))
        mapa["{{PTOT_TOT}}"]  = str(int(prod_df["total"].sum()))
        mapa["{{PTOT_H}}"]    = f"{prod_df['horas'].sum():.1f}"
        mapa["{{PTOT_MH}}"]   = f"{prod_df['media_horas'].mean():.1f}"
        mapa["{{PTOT_PUB}}"]  = str(int(prod_df["publicacoes"].sum()))
    else:
        for k in ("{{PTOT_CONC}}","{{PTOT_PEND}}","{{PTOT_TOT}}","{{PTOT_H}}","{{PTOT_MH}}","{{PTOT_PUB}}"):
            mapa[k] = "—"

    # Fluxo mensal 1–6
    for i in range(6):
        n = i + 1
        mapa[f"{{{{MES_{n}}}}}"]     = _fl(i, "mes_str")
        mapa[f"{{{{FL{n}_ENT}}}}"]   = _fl(i, "ent")
        mapa[f"{{{{FL{n}_CONC}}}}"]  = _fl(i, "conc")
        mapa[f"{{{{FL{n}_SALD}}}}"]  = _fl(i, "saldo")
        mapa[f"{{{{FL{n}_SACUM}}}}"] = _fl(i, "sacum")
        mapa[f"{{{{FL{n}_PEND}}}}"]  = _fl(i, "pend")

    # Totais fluxo
    if not flux_df.empty:
        mapa["{{FLTOT_ENT}}"]  = str(int(flux_df["ent"].sum()))
        mapa["{{FLTOT_CONC}}"] = str(int(flux_df["conc"].sum()))
        mapa["{{FLTOT_SALD}}"] = str(int(flux_df["saldo"].sum()))
        mapa["{{FLTOT_PEND}}"] = str(int(flux_df["pend"].sum()))
    else:
        for k in ("{{FLTOT_ENT}}","{{FLTOT_CONC}}","{{FLTOT_SALD}}","{{FLTOT_PEND}}"):
            mapa[k] = "—"

    # Tipos de demanda 1–8
    for i in range(8):
        n = i + 1
        mapa[f"{{{{TIPO_{n}}}}}"] = _td(i, "demanda")
        mapa[f"{{{{TD{n}_N}}}}"]  = _td(i, "n")
        mapa[f"{{{{TD{n}_P}}}}"]  = _td(i, "pend")

    # Sazonalidade (12 meses)
    for i, abrev in enumerate(_meses_abrev):
        idx = i + 1
        row = saz_agg.loc[idx] if idx in saz_agg.index else None
        mapa[f"{{{{SAZ_{abrev}_ENT}}}}"] = str(int(row["ent"]))   if row is not None else "0"
        mapa[f"{{{{SAZ_{abrev}_H}}}}"]   = f"{row['horas']:.1f}" if row is not None else "0.0"

    # Gargalos 1–10
    for i in range(10):
        n = i + 1
        mapa[f"{{{{GARG{n}_PROC}}}}"] = _g(i, "processo_base")
        mapa[f"{{{{GARG{n}_PCR}}}}"]  = _g(i, "procurador")
        mapa[f"{{{{GARG{n}_TIPO}}}}"] = _g(i, "demanda")
        mapa[f"{{{{GARG{n}_ENT}}}}"]  = _g(i, "entrada", "date")
        mapa[f"{{{{GARG{n}_DIAS}}}}"] = _g(i, "dias")
        mapa[f"{{{{GARG{n}_MAT}}}}"]  = _g(i, "materia")

    # Histórico 12 meses
    for i in range(12):
        n = i + 1
        mapa[f"{{{{HIST_MES_{n}}}}}"]  = _h(i, "mes_str")
        mapa[f"{{{{HIST{n}_ENT}}}}"]   = _h(i, "ent",  "int")
        mapa[f"{{{{HIST{n}_CONC}}}}"]  = _h(i, "conc", "int")
        mapa[f"{{{{HIST{n}_H}}}}"]     = _h(i, "horas","f1")
        mapa[f"{{{{HIST{n}_MH}}}}"]    = _h(i, "mh",   "f1")
        mapa[f"{{{{HIST{n}_PUB}}}}"]   = _h(i, "pub",  "int")
        mapa[f"{{{{HIST{n}_TX}}}}"]    = _h(i, "tx",   "f1")

    # Totais histórico
    if not hist_df.empty:
        mapa["{{HTTOT_ENT}}"]  = str(int(hist_df["ent"].sum()))
        mapa["{{HTTOT_CONC}}"] = str(int(hist_df["conc"].sum()))
        mapa["{{HTTOT_H}}"]    = f"{hist_df['horas'].sum():.1f}"
        mapa["{{HTTOT_MH}}"]   = f"{hist_df['mh'].mean():.1f}"
        mapa["{{HTTOT_PUB}}"]  = str(int(hist_df["pub"].sum()))
        tot_ent_h = int(hist_df["ent"].sum() + hist_df["conc"].sum())
        mapa["{{HTTOT_TX}}"]   = (
            f"{hist_df['conc'].sum() / tot_ent_h * 100:.1f}"
            if tot_ent_h > 0 else "—"
        )
    else:
        for k in ("{{HTTOT_ENT}}","{{HTTOT_CONC}}","{{HTTOT_H}}",
                  "{{HTTOT_MH}}","{{HTTOT_PUB}}","{{HTTOT_TX}}"):
            mapa[k] = "—"

    return mapa


def _botao_relatorio(template_name: str, mapa: dict, label: str, filename: str) -> None:
    """Renderiza o botão de download do relatório Word."""
    template_path = Path(__file__).parent / template_name
    if not mapa:
        st.warning("Sem dados para gerar o relatório.", icon="⚠️")
        return
    col_btn, _ = st.columns([1, 3])
    with col_btn:
        dados = _gerar_relatorio(template_path, mapa)
        if dados:
            st.download_button(
                label=label,
                data=dados,
                file_name=filename,
                mime=(
                    "application/vnd.openxmlformats-officedocument"
                    ".wordprocessingml.document"
                ),
                use_container_width=True,
            )

# =============================================================================
# PONTO DE ENTRADA
# =============================================================================

def main() -> None:
    _migrar_banco()

    if not _db.db_exists():
        st.markdown(
            """
            <div class="top-bar">
              <div>
                <h1>PGE <span style="font-weight:300;opacity:.85">.quant</span></h1>
                <p>Coordenadoria do Contencioso Geral &nbsp;·&nbsp; PGE/SP</p>
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        st.info(
            "**Base de dados não encontrada.** "
            "Use a aba **Atualização de Dados** para carregar os arquivos iniciais.",
            icon="🗄️",
        )
        _tab_upload()
        return

    df_full = _carregar_df()
    df_proc_filtered, filtros_sidebar = _sidebar(df_full)
    df = df_proc_filtered  # Frente 1: filtrado por acervo e titularidade do processo

    df_dem_full = _carregar_demandas()
    # Frente 2: filtrado diretamente pelas seleções da sidebar sobre as demandas,
    # sem qualquer dependência dos processos resultantes da Frente 1.
    df_dem_filtered = (
        _filtrar_demandas_frente2(df_dem_full, filtros_sidebar)
        if not df_dem_full.empty
        else pd.DataFrame()
    )

    # ── Barra superior ──────────────────────────────────────────────────────
    nucleos_ativos = st.session_state.get("f_nucleo", [])
    if nucleos_ativos:
        subtitulo = (
            " · ".join(nucleos_ativos)
            if len(nucleos_ativos) <= 3
            else f"{len(nucleos_ativos)} núcleos selecionados"
        )
    else:
        subtitulo = "Coordenadoria do Contencioso Geral"

    badge = (
        f"🔍 {len(df):,} de {len(df_full):,} processos".replace(",", ".")
        if len(df) < len(df_full)
        else f"📋 Base completa · {len(df_full):,} processos".replace(",", ".")
    )
    st.markdown(
        f"""
        <div class="top-bar">
          <div>
            <h1>PGE <span style="font-weight:300;opacity:.85">.quant</span></h1>
            <p>Coordenadoria do Contencioso Geral &nbsp;·&nbsp; {subtitulo}</p>
          </div>
          <div class="top-badge">{badge}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # ── Tabs externas ───────────────────────────────────────────────────────
    tab_fin, tab_op, tab_admin = st.tabs([
        "📊  Análise Financeira e Resultados",
        "📉  Gestão Operacional e Produtividade",
        "⚙️  Atualização de Dados",
    ])

    with tab_fin:
        s1, s2, s3, s4, s5, s6 = st.tabs([
            "Panorama Financeiro",
            "Êxito por Tese",
            "Estágio e Recursos",
            "Mapeamento Geográfico",
            "Linha do Tempo",
            "Detalhamento",
        ])
        with s1:
            _subtab_f1_panorama(df)
        with s2:
            _subtab_f1_exito_tese(df)
        with s3:
            _subtab_f1_estagio(df, df_dem_full)
        with s4:
            _subtab_f1_mapa(df)
        with s5:
            _subtab_f1_linha_tempo(df)
        with s6:
            _subtab_f1_detalhamento(df)

        st.markdown('<hr class="divider">', unsafe_allow_html=True)
        if df.empty:
            st.warning("Sem dados de processos para gerar o relatório.", icon="⚠️")
        else:
            mapa_f1 = _build_mapa_frente1(df, filtros_sidebar)
            _botao_relatorio(
                "Esqueleto_Relatorio_Frente1.docx",
                mapa_f1,
                "📄 Gerar Relatório — Frente 1",
                f"Relatorio_F1_{date.today().strftime('%Y%m%d')}.docx",
            )

    with tab_op:
        o1, o2, o3, o4, o5, o6 = st.tabs([
            "Performance dos Núcleos",
            "Fluxo de Demandas",
            "Sazonalidade Operacional",
            "Gargalos de Pendência",
            "Procuradores",
            "Histórico e Linha do Tempo",
        ])
        with o1:
            _subtab_f2_nucleos(df, df_dem_filtered)
        with o2:
            _subtab_f2_fluxo(df_dem_filtered)
        with o3:
            _subtab_f2_sazonalidade(df_dem_filtered)
        with o4:
            _subtab_f2_gargalos(df_dem_filtered)
        with o5:
            _subtab_f2_procuradores(df_dem_filtered)
        with o6:
            _subtab_f2_timeline(df, df_dem_full)

        st.markdown('<hr class="divider">', unsafe_allow_html=True)
        if df_dem_filtered.empty:
            st.warning("Sem dados de demandas para gerar o relatório.", icon="⚠️")
        else:
            mapa_f2 = _build_mapa_frente2(df_dem_filtered, filtros_sidebar)
            _botao_relatorio(
                "Esqueleto_Relatorio_Frente2.docx",
                mapa_f2,
                "📄 Gerar Relatório — Frente 2",
                f"Relatorio_F2_{date.today().strftime('%Y%m%d')}.docx",
            )

    with tab_admin:
        _tab_upload()


main()
